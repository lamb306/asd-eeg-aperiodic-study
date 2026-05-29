#!/usr/bin/env python
"""Generate NC narrative figures and insert into manuscript_NC_revised_zh_clean.docx."""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from src.config import load_config, setup_logging
from src.nc_figures import NC_OUT, generate_all_nc_figures

SRC_DOC = PROJECT_ROOT / "outputs" / "reports" / "manuscript_NC_revised_zh_clean.docx"
OUT_DOC = PROJECT_ROOT / "outputs" / "reports" / "manuscript_NC_revised_zh_with_figures.docx"
FIG_DIR = NC_OUT

FIGURE_SPECS = [
    {
        "key": "Figure1",
        "after_idx": 39,
        "title": "Figure 1. Spectral parameterization identifies lower aperiodic exponent in autistic children.",
        "caption": (
            "A, Sample inclusion flow. Among 168 children with preprocessed resting-state EEG, 145 met the "
            "minimum usable epoch criterion and 138 passed subject-level specparam quality control, yielding "
            "the primary cohort of 61 autistic children and 77 typically developing children. B, Group-averaged "
            "power spectral density showing the broadband 1/f-like background and visible alpha-range peak. "
            "C, Schematic illustration of spectral parameterization separating the aperiodic background from "
            "periodic peaks. D, Global aperiodic exponent was significantly higher in TD than ASD after "
            "adjustment for age, sex, IQ total, and usable epochs, indicating a flatter aperiodic background "
            "in ASD. E, Global offset showed a nonsignificant TD > ASD trend. F, Periodic peak parameters "
            "did not show significant group differences. Error bars indicate 95% confidence intervals where applicable."
        ),
    },
    {
        "key": "Figure2",
        "after_idx": 42,
        "title": "Figure 2. ASD-related aperiodic flattening is concentrated over posterior electrodes.",
        "caption": (
            "A, Channel-wise TD − ASD regression coefficients for aperiodic exponent across the HydroCel-64 "
            "montage. Positive values indicate higher exponent in TD than ASD. Channels surviving FDR "
            "correction are marked. B, Significant channels were clustered over posterior parietal–occipital "
            "electrodes, including E33, E36, E37, and E38. These channels were averaged to define the posterior "
            "exponent. C, Posterior exponent showed a TD > ASD pattern consistent with the global effect. "
            "D, Comparison of global and posterior effect estimates indicates that the posterior exponent "
            "captured a spatially concentrated ASD-related aperiodic alteration. Topographic results are limited "
            "to scalp electrode space and should not be interpreted as source localization."
        ),
    },
    {
        "key": "Figure3",
        "after_idx": 46,
        "title": "Figure 3. Aperiodic alterations are developmentally moderated.",
        "caption": (
            "A, Model-predicted global aperiodic exponent as a function of age by group. The group × age "
            "interaction was significant, indicating that ASD–TD differences in exponent varied with age. "
            "B, Global offset also showed a significant group × age interaction. C, Age-stratified analyses "
            "showed that the exponent group effect was not significant in children ≤72 months but was "
            "significant in children >72 months. D, Age tertile analyses further indicated that the largest "
            "TD > ASD exponent effect occurred in the oldest tertile. E, Posterior exponent also showed "
            "significant group × age moderation. Because the study is cross-sectional, these age effects should "
            "not be interpreted as within-individual longitudinal trajectories."
        ),
    },
    {
        "key": "Figure4",
        "after_idx": 51,
        "title": "Figure 4. Autistic children show negative deviations from TD normative aperiodic trajectories.",
        "caption": (
            "A, TD-only age-normative model for global aperiodic exponent with ASD observations overlaid. "
            "B, Deviation scores quantify each child's observed exponent relative to age-expected TD values. "
            "C, Posterior exponent deviation scores test whether posterior aperiodic flattening represents a "
            "stronger deviation from the TD normative trajectory. D, Association between deviation scores and "
            "age. This analysis places ASD–TD differences in a developmental reference frame but does not "
            "establish individual longitudinal trajectories."
        ),
    },
    {
        "key": "Figure5",
        "after_idx": 56,
        "title": "Figure 5. Posterior exponent effects are robust to oscillatory, IAF, and data-quality explanations.",
        "caption": (
            "A, Global aperiodic exponent group effects across nested covariate models. B, Frequency-range and "
            "aperiodic-mode sensitivity analyses showed directionally consistent TD > ASD effects, with stronger "
            "evidence under fixed-mode parameterization. C, Automated ICA sensitivity analyses showed that the "
            "global exponent effect weakened after fully automated artifact control, whereas posterior exponent "
            "remained significant across thresholds and age subsets after BH-FDR correction. D, Split-half "
            "reliability analyses indicated high within-session consistency for global exponent, global offset, "
            "and alpha peak power. E, IAF or alpha peak frequency sensitivity analysis showing alpha_cf age "
            "effects, null group effects, and preserved posterior exponent effects after alpha_cf adjustment."
        ),
    },
    {
        "key": "ExtendedData1",
        "after_idx": 63,
        "title": (
            "Extended Data Figure 1. Exploratory machine-learning classification based on periodic and "
            "aperiodic features."
        ),
        "caption": (
            "Nested cross-validation showed that aperiodic features outperformed periodic peak features in the "
            "full sample, although performance was modest. In the >72-month subset, adding age interaction "
            "terms improved performance. These results are exploratory and should not be interpreted as "
            "diagnostic-level classification."
        ),
    },
]


def add_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_figure_block(after_paragraph: Paragraph, image_path: Path, title: str, caption: str) -> None:
    spacer = add_paragraph_after(after_paragraph)
    spacer.paragraph_format.space_after = Pt(6)

    img_p = add_paragraph_after(spacer)
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(str(image_path), width=Cm(16.5))

    title_p = add_paragraph_after(img_p)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(10)

    cap_p = add_paragraph_after(title_p)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap_p.add_run(caption).font.size = Pt(9)


def insert_figures_into_docx(paths: dict[str, Path]) -> None:
    doc = Document(str(SRC_DOC))
    # Insert from bottom to top so indices stay valid
    for spec in sorted(FIGURE_SPECS, key=lambda s: s["after_idx"], reverse=True):
        img = paths.get(spec["key"])
        if img is None or not img.exists():
            raise FileNotFoundError(f"Missing figure for {spec['key']}: {img}")
        anchor = doc.paragraphs[spec["after_idx"]]
        insert_figure_block(anchor, img, spec["title"], spec["caption"])

    OUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOC))


def main() -> None:
    cfg = load_config()
    log = setup_logging(cfg, name="nc_figures_and_manuscript")
    deriv = Path(cfg["paths"]["derivatives_root"])
    outputs = Path(cfg["paths"]["outputs_root"])

    log.info("Generating NC figures...")
    paths = generate_all_nc_figures(cfg, deriv, outputs)
    for name, path in paths.items():
        log.info("  %s -> %s", name, path)

    log.info("Inserting figures into %s", OUT_DOC.name)
    shutil.copy2(SRC_DOC, OUT_DOC.with_suffix(".docx.bak"))
    insert_figures_into_docx(paths)
    log.info("Done: %s", OUT_DOC.resolve())


if __name__ == "__main__":
    main()
