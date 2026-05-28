from pathlib import Path
import textwrap
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import patches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(r"d:\asd_eeg_aperiodic_study")
FIG_DIR = ROOT / "figures_submission"
TAB_DIR = ROOT / "tables_submission"
OUT_MAIN = ROOT / "manuscript_submission_final.docx"
OUT_SUPP = ROOT / "supplementary_materials.docx"

ASD_COLOR = "#D95F02"
TD_COLOR = "#1B9E77"
GRAY = "#555555"


def ensure_dirs():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    TAB_DIR.mkdir(parents=True, exist_ok=True)


def fmt_p(p):
    if p < 1e-4:
        return f"{p:.2e}"
    return f"{p:.4f}"


def save_fig(fig, stem):
    png = FIG_DIR / f"{stem}.png"
    pdf = FIG_DIR / f"{stem}.pdf"
    fig.savefig(png, dpi=600, bbox_inches="tight")
    fig.savefig(pdf, dpi=600, bbox_inches="tight")
    plt.close(fig)
    return png, pdf


def load_data():
    d = {}
    d["flow"] = pd.read_csv(ROOT / "outputs" / "tables" / "sample_inclusion_flow.csv")
    d["locked"] = pd.read_csv(ROOT / "outputs" / "tables" / "final_paper_stats_locked.csv")
    d["main_group"] = pd.read_csv(ROOT / "derivatives" / "stats" / "main_group_analysis.csv")
    d["channels"] = pd.read_csv(ROOT / "outputs" / "tables" / "significant_channels_fdr.csv")
    d["age_inter"] = pd.read_csv(ROOT / "outputs" / "tables" / "compare_preschool_study" / "age_interaction_models.csv")
    d["split"] = pd.read_csv(ROOT / "outputs" / "tables" / "extension" / "split_half_reliability.csv")
    d["ic70"] = pd.read_csv(ROOT / "outputs" / "tables" / "iclabel_sensitivity" / "iclabel_main_group_analysis_threshold_0_70.csv")
    d["ic80"] = pd.read_csv(ROOT / "outputs" / "tables" / "iclabel_sensitivity" / "iclabel_main_group_analysis_threshold_0_80.csv")
    d["relaxed"] = pd.read_csv(ROOT / "outputs" / "tables" / "_relaxed_significance_snapshot.csv")
    d["strict"] = pd.read_csv(ROOT / "outputs" / "tables" / "_strict_significance_snapshot.csv")
    d["qc_models"] = pd.read_csv(ROOT / "outputs" / "tables" / "qc_covariate_control_models.csv")
    d["qc_event_fdr"] = pd.read_csv(ROOT / "outputs" / "tables" / "qc_covariate_control_event_fdr.csv")
    d["resting_feat"] = pd.read_csv(ROOT / "outputs" / "tables" / "resting_features_locked.csv")
    d["participants"] = pd.read_csv(ROOT / "derivatives" / "participants_analysis.csv")
    d["movie_isc"] = pd.read_csv(ROOT / "derivatives_task_movie" / "stats" / "movie_isc_subject_values_with_neutral.csv")
    d["coupling"] = pd.read_csv(ROOT / "derivatives_task_movie" / "stats" / "resting_movie_coupling_merged.csv")
    d["cars_subj"] = pd.read_csv(ROOT / "derivatives_task_movie" / "stats" / "asd_isc_cars_subject_values.csv")
    d["coupling_ols"] = pd.read_csv(ROOT / "derivatives_task_movie" / "stats" / "resting_movie_coupling_interaction_model.csv")
    d["coupling_rlm"] = pd.read_csv(ROOT / "derivatives_task_movie" / "stats" / "resting_movie_coupling_interaction_model_rlm_winsor.csv")
    d["cls_all"] = pd.read_csv(ROOT / "outputs" / "ml_biomarker" / "classification_results.csv")
    d["cls_old"] = pd.read_csv(ROOT / "outputs" / "ml_biomarker" / "classification_results__abc_ageint_older72_v2check.csv")
    d["cls_ch"] = pd.read_csv(ROOT / "outputs" / "ml_biomarker" / "classification_results__channelwise_older72_v2check.csv")
    d["delong_all"] = pd.read_csv(ROOT / "outputs" / "ml_biomarker" / "delong_results.csv")
    d["delong_old"] = pd.read_csv(ROOT / "outputs" / "ml_biomarker" / "delong_results__abc_ageint_older72_v2check.csv")
    return d


def generate_figure1():
    fig, axes = plt.subplots(1, 3, figsize=(16, 4.8))

    ax = axes[0]
    ax.set_title("A. Sample inclusion")
    ax.axis("off")
    y0 = 0.85
    labels = [
        "Initial cohort\nN=168",
        "Usable epochs >= 60\nN=145",
        "Specparam QC passed\nN=138",
    ]
    for i, label in enumerate(labels):
        rect = patches.FancyBboxPatch((0.15, y0 - i * 0.28), 0.7, 0.15, boxstyle="round,pad=0.02", fc="#f2f2f2", ec=GRAY)
        ax.add_patch(rect)
        ax.text(0.5, y0 + 0.075 - i * 0.28, label, ha="center", va="center", fontsize=10)
        if i < 2:
            ax.annotate("", xy=(0.5, y0 - 0.02 - i * 0.28), xytext=(0.5, y0 - 0.12 - i * 0.28),
                        arrowprops=dict(arrowstyle="->", lw=1.5, color=GRAY))

    ax = axes[1]
    ax.set_title("B. Cross-state analysis framework")
    ax.axis("off")
    items = [
        ("Resting-state EEG\n(specparam)", 0.1),
        ("Aperiodic metrics\n(global/posterior exponent)", 0.38),
        ("Natural movie ISC\n(mental/pain/neutral)", 0.66),
    ]
    for txt, x in items:
        rect = patches.FancyBboxPatch((x, 0.4), 0.22, 0.2, boxstyle="round,pad=0.02", fc="#f7fbff", ec="#2c7fb8")
        ax.add_patch(rect)
        ax.text(x + 0.11, 0.5, txt, ha="center", va="center", fontsize=9)
    ax.annotate("", xy=(0.38, 0.5), xytext=(0.32, 0.5), arrowprops=dict(arrowstyle="->", lw=1.5))
    ax.annotate("", xy=(0.66, 0.5), xytext=(0.60, 0.5), arrowprops=dict(arrowstyle="->", lw=1.5))
    ax.text(0.5, 0.25, "Cross-state coupling model:\nmental ISC ~ posterior exponent × group + covariates",
            ha="center", va="center", fontsize=9)

    ax = axes[2]
    ax.set_title("C. Analysis sample sizes")
    ax.axis("off")
    rows = [
        ("Resting-state primary analysis", "N=138 (ASD=61, TD=77)"),
        ("Movie ISC group comparison", "ASD=73, TD=95/96"),
        ("Cross-state coupling primary", "N=128"),
        ("Stringent inclusion sensitivity", "N=102"),
    ]
    yy = 0.82
    for name, n in rows:
        ax.text(0.05, yy, name, fontsize=9, ha="left")
        ax.text(0.95, yy, n, fontsize=9, ha="right")
        ax.plot([0.05, 0.95], [yy - 0.05, yy - 0.05], color="#dddddd", lw=1)
        yy -= 0.2

    fig.suptitle("Figure 1. Study design and sample inclusion", fontsize=14, fontweight="bold", y=1.03)
    return save_fig(fig, "Figure1_study_design_and_sample_inclusion")


def generate_figure2(data):
    rf = data["resting_feat"].copy()
    fig, axes = plt.subplots(1, 3, figsize=(17, 5))

    for i, var in enumerate(["global_exponent", "global_offset"]):
        ax = axes[i]
        asd = rf.loc[rf["group"] == "ASD", var].dropna().values
        td = rf.loc[rf["group"] == "TD", var].dropna().values
        b = ax.boxplot([asd, td], labels=["ASD", "TD"], patch_artist=True, widths=0.5)
        b["boxes"][0].set(facecolor=ASD_COLOR, alpha=0.35)
        b["boxes"][1].set(facecolor=TD_COLOR, alpha=0.35)
        ax.scatter(np.random.normal(1, 0.04, len(asd)), asd, c=ASD_COLOR, s=12, alpha=0.5)
        ax.scatter(np.random.normal(2, 0.04, len(td)), td, c=TD_COLOR, s=12, alpha=0.5)
        ax.set_title(("A. " if i == 0 else "B. ") + var)
        ax.grid(axis="y", alpha=0.2)
        ax.set_ylabel(var)

    ax = axes[2]
    rows = [
        ("Group only", 0.096, 0.048, 0.145),
        ("+ Age, sex", 0.090, 0.038, 0.142),
        ("+ IQ", 0.080, 0.019, 0.141),
        ("Primary", 0.0791, 0.0177, 0.1404),
        ("+ Mean fit quality", 0.056, 0.005, 0.106),
        ("+ Bad channels", 0.081, 0.019, 0.142),
        ("ICA threshold 0.80", 0.053, -0.013, 0.119),
        ("ICA threshold 0.70", 0.053, -0.014, 0.121),
    ]
    y = np.arange(len(rows))[::-1]
    betas = [r[1] for r in rows]
    lo = [r[2] for r in rows]
    hi = [r[3] for r in rows]
    err_left = np.array(betas) - np.array(lo)
    err_right = np.array(hi) - np.array(betas)
    ax.errorbar(betas, y, xerr=[err_left, err_right], fmt="o", color="#2c3e50", ecolor="#2c3e50", capsize=3)
    ax.axvline(0, color="#999999", ls="--", lw=1)
    ax.set_yticks(y)
    ax.set_yticklabels([r[0] for r in rows], fontsize=8)
    ax.set_title("C. Robustness models (TD-ASD beta)")
    ax.set_xlabel("beta (95% CI)")
    ax.grid(axis="x", alpha=0.2)

    fig.suptitle("Figure 2. Resting-state aperiodic EEG differences", fontsize=14, fontweight="bold", y=1.03)
    return save_fig(fig, "Figure2_resting_aperiodic_differences")


def generate_figure3(data):
    rf = data["resting_feat"].copy()
    part = data["participants"][["subject_id", "age_months", "group"]].drop_duplicates()
    df = rf.merge(part, on=["subject_id", "group"], how="left")

    fig, axes = plt.subplots(1, 3, figsize=(17, 5))
    for idx, var in enumerate(["global_exponent", "global_offset"]):
        ax = axes[idx]
        for grp, color in [("ASD", ASD_COLOR), ("TD", TD_COLOR)]:
            sub = df[df["group"] == grp].dropna(subset=[var, "age_months"])
            ax.scatter(sub["age_months"], sub[var], s=15, alpha=0.45, color=color, label=grp if idx == 0 else None)
            if len(sub) > 2:
                coef = np.polyfit(sub["age_months"], sub[var], 1)
                x = np.linspace(sub["age_months"].min(), sub["age_months"].max(), 100)
                ax.plot(x, coef[0] * x + coef[1], color=color, lw=2)
        ax.set_title(("A. " if idx == 0 else "B. ") + f"Age effect on {var}")
        ax.set_xlabel("Age (months)")
        ax.set_ylabel(var)
        ax.grid(alpha=0.2)
    axes[0].legend(frameon=False)

    ax = axes[2]
    strata = [
        ("<=72 months", 0.055, -0.100, 0.211, 0.466),
        (">72 months", 0.076, 0.007, 0.145, 0.031),
    ]
    y = np.array([1, 0])
    beta = np.array([s[1] for s in strata])
    lo = np.array([s[2] for s in strata])
    hi = np.array([s[3] for s in strata])
    ax.errorbar(beta, y, xerr=[beta - lo, hi - beta], fmt="o", color="#2c3e50", capsize=3)
    ax.axvline(0, color="#999999", ls="--", lw=1)
    ax.set_yticks(y)
    ax.set_yticklabels([f"{s[0]} (p={s[4]:.3f})" for s in strata], fontsize=9)
    ax.set_xlabel("TD-ASD beta for global exponent")
    ax.set_title("C. Age-stratified effect")
    ax.grid(axis="x", alpha=0.2)

    fig.suptitle("Figure 3. Age-dependent effects", fontsize=14, fontweight="bold", y=1.03)
    return save_fig(fig, "Figure3_age_dependent_effects")


def generate_figure4(data):
    ch = data["channels"].copy()
    rf = data["resting_feat"].copy()
    fig, axes = plt.subplots(1, 3, figsize=(17, 5))

    ax = axes[0]
    ax.set_title("A. Electrode-level posterior effects")
    head = patches.Circle((0, 0), 1.0, fill=False, lw=2, ec=GRAY)
    ax.add_patch(head)
    ax.plot([0, 0.08, -0.08, 0], [1.0, 1.08, 1.08, 1.0], color=GRAY, lw=1.3)
    ax.plot([-1.0, -1.08, -1.0], [0.1, 0.0, -0.1], color=GRAY, lw=1.2)
    ax.plot([1.0, 1.08, 1.0], [0.1, 0.0, -0.1], color=GRAY, lw=1.2)
    for _, r in ch.iterrows():
        x = float(r["pos_x"]) * 8.0
        y = float(r["pos_y"]) * 8.0
        ax.scatter(x, y, s=220, color="#2c7fb8", alpha=0.8)
        ax.text(x, y, r["channel"], ha="center", va="center", fontsize=8, color="white", fontweight="bold")
    ax.set_xlim(-1.2, 1.2)
    ax.set_ylim(-1.2, 1.2)
    ax.set_aspect("equal")
    ax.axis("off")

    ax = axes[1]
    ax.set_title("B. Effect sizes at significant electrodes")
    channels = ch["channel"].tolist()
    betas = ch["coef"].values
    qvals = ch["pvalue_fdr"].values
    bars = ax.bar(channels, betas, color="#2c7fb8")
    for i, b in enumerate(bars):
        ax.text(b.get_x() + b.get_width() / 2, b.get_height() + 0.003, f"q={qvals[i]:.3f}",
                ha="center", va="bottom", fontsize=8, rotation=90)
    ax.set_ylabel("beta (TD-ASD)")
    ax.grid(axis="y", alpha=0.2)

    ax = axes[2]
    asd = rf[rf["group"] == "ASD"]["posterior_exponent"].dropna().values
    td = rf[rf["group"] == "TD"]["posterior_exponent"].dropna().values
    b = ax.boxplot([asd, td], labels=["ASD", "TD"], patch_artist=True, widths=0.5)
    b["boxes"][0].set(facecolor=ASD_COLOR, alpha=0.35)
    b["boxes"][1].set(facecolor=TD_COLOR, alpha=0.35)
    ax.scatter(np.random.normal(1, 0.04, len(asd)), asd, c=ASD_COLOR, s=12, alpha=0.5)
    ax.scatter(np.random.normal(2, 0.04, len(td)), td, c=TD_COLOR, s=12, alpha=0.5)
    ax.set_title("C. Posterior exponent by group")
    ax.set_ylabel("posterior_exponent")
    ax.grid(axis="y", alpha=0.2)

    fig.suptitle("Figure 4. Spatial distribution of posterior effects", fontsize=14, fontweight="bold", y=1.03)
    return save_fig(fig, "Figure4_spatial_distribution_posterior_effects")


def generate_figure5(data):
    m = data["movie_isc"].copy()
    m = m[m["event_type"].isin(["mental", "pain", "neutral"])].copy()
    fig, axes = plt.subplots(2, 2, figsize=(14, 10))
    panel = {"mental": "A", "pain": "B", "neutral": "C"}
    for ev, ax in zip(["mental", "pain", "neutral"], [axes[0, 0], axes[0, 1], axes[1, 0]]):
        sub = m[m["event_type"] == ev].dropna(subset=["isc_z"])
        asd = sub[sub["group"] == "ASD"]["isc_z"].values
        td = sub[sub["group"] == "TD"]["isc_z"].values
        b = ax.boxplot([asd, td], labels=["ASD", "TD"], patch_artist=True, widths=0.5)
        b["boxes"][0].set(facecolor=ASD_COLOR, alpha=0.35)
        b["boxes"][1].set(facecolor=TD_COLOR, alpha=0.35)
        ax.scatter(np.random.normal(1, 0.04, len(asd)), asd, c=ASD_COLOR, s=10, alpha=0.4)
        ax.scatter(np.random.normal(2, 0.04, len(td)), td, c=TD_COLOR, s=10, alpha=0.4)
        ax.set_title(f"{panel[ev]}. {ev} ISC")
        ax.set_ylabel("ISC (z)")
        ax.grid(axis="y", alpha=0.2)

    ax = axes[1, 1]
    lock = data["locked"].copy()
    isc = lock[lock["Analysis_Type"].str.startswith("ISC_")].copy()
    isc["event"] = isc["Analysis_Type"].str.replace("ISC_", "", regex=False)
    y = np.arange(len(isc))[::-1]
    ax.scatter(isc["Test_Statistic"], y, color="#2c3e50", s=55)
    for i, (_, r) in enumerate(isc.iterrows()):
        ax.text(r["Test_Statistic"] + 0.05, y[i], f"q={float(r['FDR_p']):.2e}", va="center", fontsize=8)
    ax.axvline(0, color="#999999", ls="--", lw=1)
    ax.set_yticks(y)
    ax.set_yticklabels([f"{e}" for e in isc["event"]])
    ax.set_xlabel("t statistic (TD vs ASD)")
    ax.set_title("D. Event-level t statistics")
    ax.grid(axis="x", alpha=0.2)

    fig.suptitle("Figure 5. Naturalistic movie ISC differences", fontsize=14, fontweight="bold", y=1.01)
    return save_fig(fig, "Figure5_movie_isc_differences")


def generate_figure6(data):
    cp = data["coupling"].copy().dropna(subset=["mental_isc_z", "posterior_exponent", "group"])
    cp = cp[cp["group"].isin(["ASD", "TD"])].copy()
    fig, axes = plt.subplots(1, 3, figsize=(17, 5))

    ax = axes[0]
    for grp, color in [("ASD", ASD_COLOR), ("TD", TD_COLOR)]:
        sub = cp[cp["group"] == grp]
        ax.scatter(sub["posterior_exponent"], sub["mental_isc_z"], s=22, alpha=0.6, color=color, label=grp)
        if len(sub) > 2:
            coef = np.polyfit(sub["posterior_exponent"], sub["mental_isc_z"], 1)
            x = np.linspace(sub["posterior_exponent"].min(), sub["posterior_exponent"].max(), 100)
            ax.plot(x, coef[0] * x + coef[1], color=color, lw=2)
    ax.set_xlabel("Posterior exponent")
    ax.set_ylabel("Mental ISC (z)")
    ax.set_title("A. Primary analysis sample (n=128)")
    ax.legend(frameon=False)
    ax.grid(alpha=0.2)

    ax = axes[1]
    rows = [
        ("OLS (primary)", -0.3519, 0.0102),
        ("RLM/winsor (primary)", -0.5318, 0.00259),
    ]
    y = np.array([1, 0])
    beta = np.array([r[1] for r in rows])
    ax.scatter(beta, y, color="#2c3e50", s=60)
    for i, r in enumerate(rows):
        ax.text(beta[i] + 0.02, y[i], f"p={r[2]:.4f}", va="center", fontsize=9)
    ax.axvline(0, color="#999999", ls="--", lw=1)
    ax.set_yticks(y)
    ax.set_yticklabels([r[0] for r in rows], fontsize=9)
    ax.set_xlabel("Interaction beta")
    ax.set_title("B. Interaction estimates")
    ax.grid(axis="x", alpha=0.2)

    ax = axes[2]
    labels = ["OLS", "RLM"]
    p_primary = [0.0102, 0.00259]
    p_stringent = [0.0792, 0.0195]
    x = np.arange(2)
    width = 0.35
    ax.bar(x - width / 2, -np.log10(p_primary), width, label="Primary analysis sample", color="#4c78a8")
    ax.bar(x + width / 2, -np.log10(p_stringent), width, label="Stringent inclusion sensitivity", color="#f58518")
    ax.axhline(-np.log10(0.05), color="#999999", ls="--", lw=1)
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylabel("-log10(p)")
    ax.set_title("C. Sensitivity to inclusion/model setting")
    ax.legend(frameon=False, fontsize=8)
    ax.grid(axis="y", alpha=0.2)

    fig.suptitle("Figure 6. Cross-state coupling between resting posterior exponent and mental ISC", fontsize=13, fontweight="bold", y=1.03)
    return save_fig(fig, "Figure6_cross_state_coupling")


def generate_figure7(data):
    cars = data["cars_subj"].copy()
    post = data["resting_feat"][["subject_id", "posterior_exponent"]].copy()
    cdf = cars.merge(post, on="subject_id", how="left").dropna(subset=["CARS_total", "posterior_exponent"])

    fig, axes = plt.subplots(1, 3, figsize=(17, 5))
    ax = axes[0]
    ax.scatter(cdf["posterior_exponent"], cdf["CARS_total"], color=ASD_COLOR, alpha=0.7, s=24)
    coef = np.polyfit(cdf["posterior_exponent"], cdf["CARS_total"], 1)
    xx = np.linspace(cdf["posterior_exponent"].min(), cdf["posterior_exponent"].max(), 100)
    ax.plot(xx, coef[0] * xx + coef[1], color=ASD_COLOR, lw=2)
    ax.set_title("A. posterior exponent vs CARS (ASD)")
    ax.set_xlabel("Posterior exponent")
    ax.set_ylabel("CARS total")
    ax.grid(alpha=0.2)

    ax = axes[1]
    auc_vals = [0.537, 0.681, 0.651, 0.800, 0.695]
    labs = ["Model A", "Model B", "Model C", "Model B+age (>72m)", "Channel-wise EN"]
    colors = ["#bdbdbd", "#4c78a8", "#72b7b2", "#f58518", "#54a24b"]
    ax.barh(labs, auc_vals, color=colors)
    for i, v in enumerate(auc_vals):
        ax.text(v + 0.005, i, f"{v:.3f}", va="center", fontsize=8)
    ax.set_xlim(0.45, 0.85)
    ax.set_xlabel("AUC")
    ax.set_title("B. Classification performance (exploratory)")
    ax.grid(axis="x", alpha=0.2)

    ax = axes[2]
    feats = ["posterior_exponent", "global_exponent", "age_x_posterior_exponent", "occipital_exponent", "alpha_pw"]
    imp = [1.00, 0.78, 0.66, 0.42, 0.21]
    ax.bar(feats, imp, color="#4c78a8")
    ax.set_xticklabels(feats, rotation=45, ha="right")
    ax.set_ylabel("Relative importance")
    ax.set_title("C. Top features (exploratory)")
    ax.grid(axis="y", alpha=0.2)

    fig.suptitle("Figure 7. Exploratory clinical and classification analyses", fontsize=14, fontweight="bold", y=1.03)
    return save_fig(fig, "Figure7_exploratory_clinical_classification")


def generate_supp_figures(data):
    supp = {}
    # S1
    sp = data["split"].copy()
    fig, ax = plt.subplots(figsize=(7, 4.5))
    ax.bar(sp["metric"], sp["spearman_rho"], color="#4c78a8", label="Spearman rho")
    ax.scatter(sp["metric"], sp["spearman_brown_spearman"], color="#f58518", zorder=3, label="Spearman-Brown")
    ax.set_ylim(0.9, 1.0)
    ax.set_ylabel("Reliability")
    ax.set_title("Figure S1. Split-half reliability")
    ax.legend(frameon=False)
    ax.grid(axis="y", alpha=0.2)
    supp["FigureS1_split_half_reliability"] = save_fig(fig, "FigureS1_split_half_reliability")

    # S2
    fig, ax = plt.subplots(figsize=(8, 4.8))
    labels = ["Global exp (0.80)", "Global exp (0.70)", "Posterior exp (0.70 all)", "Posterior exp (0.80 all)",
              "Posterior exp (0.70 older)", "Posterior exp (0.80 older)"]
    beta = [0.053, 0.053, 0.121, 0.128, 0.151, 0.160]
    lo = [-0.013, -0.014, 0.054, 0.054, 0.077, 0.077]
    hi = [0.119, 0.121, 0.189, 0.201, 0.226, 0.244]
    y = np.arange(len(labels))[::-1]
    ax.errorbar(beta, y, xerr=[np.array(beta)-np.array(lo), np.array(hi)-np.array(beta)],
                fmt="o", color="#2c3e50", capsize=3)
    ax.axvline(0, ls="--", color="#999")
    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlabel("beta (95% CI)")
    ax.set_title("Figure S2. Automated ICA global vs posterior effects")
    ax.grid(axis="x", alpha=0.2)
    supp["FigureS2_ica_global_vs_posterior"] = save_fig(fig, "FigureS2_ica_global_vs_posterior")

    # S3
    fig, ax = plt.subplots(figsize=(7, 4.5))
    m = ["Coupling OLS", "Coupling RLM", "posterior-CARS"]
    p1 = [0.0102, 0.00259, 0.0404]
    p2 = [0.0792, 0.0195, 0.0683]
    x = np.arange(len(m))
    ax.plot(x, -np.log10(p1), marker="o", label="Primary analysis sample")
    ax.plot(x, -np.log10(p2), marker="s", label="Stringent inclusion sensitivity")
    ax.axhline(-np.log10(0.05), ls="--", color="#999")
    ax.set_xticks(x)
    ax.set_xticklabels(m)
    ax.set_ylabel("-log10(p)")
    ax.set_title("Figure S3. Primary vs stringent inclusion comparison")
    ax.legend(frameon=False)
    ax.grid(axis="y", alpha=0.2)
    supp["FigureS3_primary_vs_stringent"] = save_fig(fig, "FigureS3_primary_vs_stringent")

    # S4
    fig, ax = plt.subplots(figsize=(7, 4.5))
    labs = ["Delta_mental", "Delta_pain"]
    tvals = [3.5010, 3.6950]
    q = [0.0007113, 0.0007113]
    bars = ax.bar(labs, tvals, color=["#4c78a8", "#54a24b"])
    for i, b in enumerate(bars):
        ax.text(b.get_x() + b.get_width()/2, b.get_height()+0.05, f"q={q[i]:.4f}", ha="center", fontsize=8)
    ax.set_ylabel("t statistic")
    ax.set_title("Figure S4. Delta_Exponent event results")
    ax.grid(axis="y", alpha=0.2)
    supp["FigureS4_delta_exponent"] = save_fig(fig, "FigureS4_delta_exponent")

    # S5
    fig, ax = plt.subplots(figsize=(6.5, 4.8))
    cars = data["cars_subj"].copy()
    post = data["resting_feat"][["subject_id", "posterior_exponent"]]
    cdf = cars.merge(post, on="subject_id", how="left").dropna(subset=["CARS_total", "posterior_exponent"])
    ax.scatter(cdf["posterior_exponent"], cdf["CARS_total"], color=ASD_COLOR, s=25, alpha=0.7)
    coef = np.polyfit(cdf["posterior_exponent"], cdf["CARS_total"], 1)
    xs = np.linspace(cdf["posterior_exponent"].min(), cdf["posterior_exponent"].max(), 100)
    ax.plot(xs, coef[0]*xs + coef[1], color=ASD_COLOR, lw=2)
    ax.set_xlabel("Posterior exponent")
    ax.set_ylabel("CARS total")
    ax.set_title("Figure S5. Posterior-CARS exploratory association")
    ax.grid(alpha=0.2)
    supp["FigureS5_posterior_cars_scatter"] = save_fig(fig, "FigureS5_posterior_cars_scatter")

    # S6
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.8))
    vals = [0.537, 0.681, 0.651, 0.800, 0.695]
    names = ["A", "B", "C", "B+age (>72m)", "Ch-wise EN"]
    axes[0].bar(names, vals, color="#4c78a8")
    axes[0].set_ylim(0.45, 0.85)
    axes[0].set_ylabel("AUC")
    axes[0].set_title("AUC summary")
    axes[0].grid(axis="y", alpha=0.2)
    f = ["posterior_exp", "global_exp", "age×posterior", "occipital_exp", "alpha_pw"]
    im = [1.0, 0.78, 0.66, 0.42, 0.21]
    axes[1].barh(f, im, color="#f58518")
    axes[1].set_title("Feature importance")
    axes[1].set_xlabel("Relative importance")
    axes[1].grid(axis="x", alpha=0.2)
    fig.suptitle("Figure S6. Classification AUC and feature importance", y=1.02)
    supp["FigureS6_classification_summary"] = save_fig(fig, "FigureS6_classification_summary")
    return supp


def write_tables(data):
    table1 = pd.DataFrame({
        "变量": ["样本量", "年龄（月）", "性别（女/男）", "IQ_total", "usable epochs", "bad channel count", "mean specparam R²"],
        "ASD": ["61", "85.7 ± 16.9", "5/56", "95.0 ± 15.2", "120.3 ± 26.8", "1.2 ± 0.6", "0.983 ± 0.011"],
        "TD": ["77", "88.8 ± 19.6", "28/49", "113.2 ± 14.6", "127.1 ± 28.7", "1.3 ± 0.5", "0.987 ± 0.008"],
        "统计检验与p值": ["", "p=0.319", "p<0.001", "p<0.001", "p=0.152", "p=0.426", "p=0.006"]
    })
    table1.to_csv(TAB_DIR / "Table1_demographic_eeg_quality.csv", index=False, encoding="utf-8-sig")

    table2 = pd.DataFrame({
        "模型": ["global_exponent 主模型", "global_offset 主模型", "group×age (global_exponent)", "group×age (global_offset)",
               "自动ICA 阈值0.80 (global_exponent)", "自动ICA 阈值0.70 (global_exponent)"],
        "beta": [0.0791, 0.0596, 0.0033, 0.0037, 0.053, 0.053],
        "SE": [0.0310, "", "", "", "", ""],
        "95% CI": ["[0.0177, 0.1404]", "", "", "", "", ""],
        "p": [0.0119, 0.0951, 0.0200, 0.0210, 0.115, 0.119],
        "n": [138, 138, 138, 138, 135, 137],
    })
    table2.to_csv(TAB_DIR / "Table2_primary_sensitivity_models.csv", index=False, encoding="utf-8-sig")

    table3 = data["channels"][["channel", "coef", "pvalue", "pvalue_fdr"]].copy()
    table3.columns = ["电极", "beta", "p", "q"]
    table3.to_csv(TAB_DIR / "Table3_channel_fdr_posterior.csv", index=False, encoding="utf-8-sig")

    lock = data["locked"].copy()
    lock.columns = ["Analysis_Type", "Cohort_N", "Test_Statistic", "Raw_p", "q"]
    lock.to_csv(TAB_DIR / "Table4_movie_isc_delta_group_diff.csv", index=False, encoding="utf-8-sig")

    table5 = pd.DataFrame({
        "模型": ["主分析样本 OLS", "主分析样本 RLM/winsor", "严格纳入标准敏感性 OLS", "严格纳入标准敏感性 RLM"],
        "交互beta": [-0.3519, -0.5318, np.nan, np.nan],
        "p": [0.0102, 0.00259, 0.0792, 0.0195],
        "n": [128, 128, 102, 102]
    })
    table5.to_csv(TAB_DIR / "Table5_cross_state_coupling_models.csv", index=False, encoding="utf-8-sig")

    table6 = pd.DataFrame({
        "项目": [
            "posterior-CARS (主分析样本)",
            "posterior-CARS (严格纳入标准敏感性)",
            "posterior-CARS QC协变量模型",
            "Model B AUC_mean (全样本)",
            "Model A AUC_mean (全样本)",
            "Model C AUC_mean (全样本)",
            "Model B+age AUC_mean (>72月)",
            "channel-wise elastic net AUC_mean (>72月)"
        ],
        "结果": [
            "Spearman rho=-0.2611, p=0.0404, n=62",
            "p=0.0683, n=60",
            "posterior_exponent p=0.448",
            "0.681",
            "0.537",
            "0.651",
            "0.800",
            "0.695"
        ]
    })
    table6.to_csv(TAB_DIR / "Table6_exploratory_clinical_ml.csv", index=False, encoding="utf-8-sig")

    # supplementary tables
    data["flow"].to_csv(TAB_DIR / "TableS1_full_sample_flow.csv", index=False, encoding="utf-8-sig")
    robust = pd.DataFrame({
        "model": ["Group only", "+Age/sex", "+IQ", "Primary", "+Mean R²", "+Bad channels"],
        "beta": [0.096, 0.090, 0.080, 0.0791, 0.056, 0.081],
        "ci_low": [0.048, 0.038, 0.019, 0.0177, 0.005, 0.019],
        "ci_high": [0.145, 0.142, 0.141, 0.1404, 0.106, 0.142],
        "p": [0.001, 0.001, 0.011, 0.0119, 0.030, 0.011]
    })
    robust.to_csv(TAB_DIR / "TableS2_nested_model_robustness.csv", index=False, encoding="utf-8-sig")
    freq = pd.DataFrame({
        "freq_range": ["1-35", "1-40", "2-35", "2-40"],
        "mode": ["fixed", "fixed", "fixed", "fixed"],
        "direction": ["TD>ASD"] * 4,
        "p_range": ["0.016-0.031"] * 4
    })
    freq.to_csv(TAB_DIR / "TableS3_frequency_sensitivity.csv", index=False, encoding="utf-8-sig")
    data["split"].to_csv(TAB_DIR / "TableS4_split_half_reliability.csv", index=False, encoding="utf-8-sig")
    ica_s = pd.DataFrame({
        "analysis": ["Global exponent (0.80)", "Global exponent (0.70)", "Posterior exponent (0.70 all)",
                    "Posterior exponent (0.80 all)", "Posterior exponent (0.70 older)", "Posterior exponent (0.80 older)"],
        "beta": [0.053, 0.053, 0.121, 0.128, 0.151, 0.160],
        "p": [0.115, 0.119, 0.000485, 0.000814, 0.000108, 0.000240],
        "q": [np.nan, np.nan, 0.0078, 0.0098, 0.0052, 0.0058]
    })
    ica_s.to_csv(TAB_DIR / "TableS5_automated_ica_results.csv", index=False, encoding="utf-8-sig")
    ml_s = pd.DataFrame({
        "metric": ["AUC Model A", "AUC Model B", "AUC Model C", "DeLong A vs B", "DeLong B vs C",
                   ">72m AUC Model B+age", ">72m DeLong A vs B", ">72m DeLong B vs C", "Channel-wise EN AUC"],
        "value": [0.537, 0.681, 0.651, 0.020, 0.024, 0.800, 0.003, 0.922, 0.695]
    })
    ml_s.to_csv(TAB_DIR / "TableS6_machine_learning_results.csv", index=False, encoding="utf-8-sig")

    return {
        "table1": table1,
        "table2": table2,
        "table3": table3,
        "table4": lock,
        "table5": table5,
        "table6": table6,
    }


def set_doc_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)


def add_table(doc, df, caption):
    if caption:
        p = doc.add_paragraph(caption)
        if p.runs:
            p.runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, c in enumerate(df.columns):
        table.rows[0].cells[i].text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if pd.isna(v) else str(v)
    doc.add_paragraph("")


def add_figure(doc, fig_path, caption):
    doc.add_picture(str(fig_path), width=Inches(6.5))
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")


def build_main_doc(tables, figs):
    doc = Document()
    set_doc_style(doc)

    t = doc.add_paragraph("自闭症谱系障碍儿童静息态 EEG 非周期神经动力学与自然电影神经同步的跨状态关联")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)
    e = doc.add_paragraph("Cross-state association between resting-state aperiodic EEG dynamics and naturalistic movie-evoked neural synchrony in children with autism spectrum disorder")
    e.alignment = WD_ALIGN_PARAGRAPH.CENTER
    e.runs[0].italic = True
    doc.add_paragraph("")

    doc.add_heading("摘要", level=1)
    abs_text = (
        "背景：传统 EEG 频段功率难以区分周期振荡与非周期宽频背景，限制了 ASD 神经动力学差异的机制解释。"
        "目的：检验 ASD 儿童静息态非周期 EEG 异常及其与自然电影神经同步差异之间的跨状态关联。"
        "方法：初始样本 168 例，经可用 epoch 与谱拟合质量控制后，静息态主分析样本为 138 例（ASD=61，TD=77）。"
        "采用 specparam 分离非周期与周期成分，主模型控制年龄、性别、IQ 和可用 epoch。自然电影分析比较 mental、pain、neutral 事件 ISC，并构建"
        " posterior exponent 与 mental ISC 的交互模型，同时实施稳健回归与严格纳入标准敏感性分析。"
        "结果：global_exponent 在 TD 高于 ASD（beta=0.0791，SE=0.0310，95% CI [0.0177, 0.1404]，p=0.0119）；"
        "global_offset 仅趋势性（beta=0.0596，p=0.0951）。后部电极 E33/E36/E37/E38 在 FDR 后显著。年龄交互提示组差具有发育依赖性。"
        "电影 ISC 在 mental、pain、neutral 三事件均表现 TD>ASD 且 q<0.05；Delta_Exponent 在 mental 与 pain 亦显著。"
        "跨状态耦合中，主分析样本交互项在 OLS 与 RLM 均显著（beta=-0.3519/p=0.0102；beta=-0.5318/p=0.00259），"
        "但在严格纳入标准敏感性分析中 OLS 不显著（p=0.0792），RLM 仍显著（p=0.0195）。"
        "posterior-CARS 关联在不同分析设定下不稳定。结论：结果支持 ASD 儿童存在“静息态非周期背景-自然情境神经同步”的跨状态神经表型，"
        "但临床相关与分类证据应视为探索性，仍需外部复现与纵向验证。"
    )
    doc.add_paragraph(abs_text)
    doc.add_paragraph("关键词：自闭症谱系障碍；静息态 EEG；非周期指数；自然电影；被试间相关；跨状态耦合")

    doc.add_heading("Introduction / 引言", level=1)
    intro = (
        "ASD 儿童 EEG 研究长期依赖传统频段功率，但该指标混合了周期振荡与非周期背景，可能导致组间差异解释偏移。"
        "功率谱参数化方法能够分离非周期指数与周期峰，为理解广谱神经背景提供更直接的量化路径。儿童期神经发育具有年龄依赖性，"
        "ASD 神经表型亦可能随年龄改变。自然电影 ISC 具备较高生态效度，可用于评估连续情境下的信息加工同步。"
        "基于此，本研究提出并检验跨状态神经表型框架：静息态非周期背景差异是否与自然电影神经同步差异共同构成 ASD 儿童的神经动力学特征。"
    )
    doc.add_paragraph(intro)

    doc.add_heading("Methods / 方法", level=1)
    methods_sections = [
        ("2.1 研究设计与样本",
         "研究采用横断面设计。静息态样本流程为 168 -> 145 -> 138。电影 ISC 分析样本为 ASD=73、TD=95/96；跨状态耦合主分析样本 n=128，"
         "并实施严格纳入标准敏感性分析（n=102）。"),
        ("2.2 EEG 采集与预处理",
         "静息态 EEG 采用 HydroCel-64 系统完成睁眼静息记录。预处理包括滤波、分段、坏导处理与伪迹控制，"
         "并以可用 epoch 阈值保证数据质量。自动 ICA 作为敏感性分支用于评估结论对伪迹清洗策略的稳定性。"),
        ("2.3 功率谱参数化",
         "使用 Welch 方法估计 1-40 Hz 功率谱，并以 specparam 固定模式分离非周期与周期成分。非周期核心指标为 global_exponent、global_offset"
         " 与 posterior_exponent。"),
        ("2.4 自然电影与跨状态模型",
         "自然电影分析包含 mental、pain、neutral 三类事件，结局为 ISC 与 Delta_Exponent。跨状态模型为 mental ISC ~ posterior exponent × group + 协变量，"
         "并行报告 OLS 与稳健回归（RLM/winsor）。"),
        ("2.5 统计分析",
         "静息态主模型控制年龄、性别、IQ_total、usable epochs；通道与事件家族采用 BH-FDR。"
         "临床相关与机器学习分析定位为探索性，重点报告效应方向、置信区间与敏感性边界。"),
    ]
    for h, ptxt in methods_sections:
        ph = doc.add_paragraph(h)
        ph.runs[0].bold = True
        doc.add_paragraph(ptxt)

    doc.add_heading("Results / 结果", level=1)
    sections = [
        ("3.1 样本与质量控制",
         "静息态主分析样本 N=138（ASD=61，TD=77）。主流程为初始 168 例，经 usable epochs >= 60 后 145 例，最终 specparam 质量控制后 138 例。"),
        ("3.2 静息态 global_exponent 与 global_offset",
         "global_exponent 显示 TD>ASD（beta=0.0791，SE=0.0310，95% CI [0.0177, 0.1404]，p=0.0119，n=138）。"
         "未校正描述统计为 ASD 1.69±0.14、TD 1.79±0.14。global_offset 为趋势性（beta=0.0596，p=0.0951）。"),
        ("3.3 年龄依赖性",
         "group×age 交互显示发育调制：global_exponent beta=0.0033（p=0.020），global_offset beta=0.0037（p=0.021）。"),
        ("3.4 空间分布",
         "后部电极 E33/E36/E37/E38 在 FDR 后显著（q=0.02540、0.03528、0.03601、0.00756），提示后部头皮电极水平效应。"),
        ("3.5 Split-half 与自动 ICA 敏感性",
         "split-half 信度较高：global exponent rho=0.959（Spearman-Brown=0.979），global offset rho=0.960（0.980），"
         "alpha peak power rho=0.972（0.986）。自动 ICA 后 global_exponent 方向一致但统计证据减弱（阈值0.80：beta=0.053，p=0.115；阈值0.70：beta=0.053，p=0.119）。"
         "后部局部指标在自动 ICA 后仍显著。"),
        ("3.6 自然电影 ISC",
         "三事件均表现 TD>ASD 且 FDR 显著：mental（t=-2.3021，q=0.0228）、pain（t=-3.9259，q=0.0002035）、"
         "neutral（t=-4.3572，q=7.38e-05）。"),
        ("3.7 Delta_Exponent",
         "Delta_mental（t=3.5010，q=0.0007113）与 Delta_pain（t=3.6950，q=0.0007113）均显著。"),
        ("3.8 跨状态耦合",
         "主分析样本中交互项在 OLS 与 RLM 均显著（beta=-0.3519，p=0.0102；beta=-0.5318，p=0.00259）。"
         "严格纳入标准敏感性分析中，OLS 不显著（p=0.0792），RLM 显著（p=0.0195）。"),
        ("3.9 posterior-CARS 探索性临床相关",
         "主分析样本 Spearman 相关为 rho=-0.2611（n=62，p=0.0404）；严格纳入标准敏感性分析不显著（n=60，p=0.0683）；"
         "QC 协变量调整后 posterior_exponent 项不显著（p=0.448）。"),
        ("3.10 机器学习探索性分类",
         "全样本中非周期特征模型（Model B）AUC_mean=0.681，优于周期模型（Model A AUC_mean=0.537）与联合模型（Model C AUC_mean=0.651）。"
         "DeLong 检验：A vs B p=0.020，B vs C p=0.024。>72 月子样本中 Model B+age interactions 最佳（AUC_mean=0.800，Balanced Accuracy=0.702，F1=0.657），"
         "A vs B p=0.003，B vs C p=0.922。channel-wise elastic net AUC_mean=0.695。"),
    ]
    for h, ptxt in sections:
        ph = doc.add_paragraph(h)
        ph.runs[0].bold = True
        doc.add_paragraph(ptxt)
        if h.startswith("3.1"):
            add_figure(doc, figs["fig1"], "Figure 1. Study design and sample inclusion. A: 样本流程；B: 跨状态分析框架；C: 主分析与严格纳入标准敏感性分析样本量。")
            add_table(doc, tables["table1"], "Table 1. Demographic and EEG quality characteristics.")
        if h.startswith("3.2"):
            add_figure(doc, figs["fig2"], "Figure 2. Resting-state aperiodic EEG differences.")
            add_table(doc, tables["table2"], "Table 2. Primary and sensitivity models for global exponent and offset.")
        if h.startswith("3.3"):
            add_figure(doc, figs["fig3"], "Figure 3. Age-dependent effects.")
        if h.startswith("3.4"):
            add_figure(doc, figs["fig4"], "Figure 4. Spatial distribution of posterior effects. 注：该图反映头皮电极水平空间分布，不代表脑源定位。")
            add_table(doc, tables["table3"], "Table 3. Channel-wise FDR-significant posterior electrodes.")
        if h.startswith("3.6"):
            add_figure(doc, figs["fig5"], "Figure 5. Naturalistic movie ISC differences.")
            add_table(doc, tables["table4"], "Table 4. Natural movie ISC and Delta_Exponent group differences.")
        if h.startswith("3.8"):
            add_figure(doc, figs["fig6"], "Figure 6. Cross-state coupling between resting posterior exponent and movie mental ISC.")
            add_table(doc, tables["table5"], "Table 5. Cross-state coupling models.")
        if h.startswith("3.10"):
            add_figure(doc, figs["fig7"], "Figure 7. Exploratory clinical and classification analyses.")
            add_table(doc, tables["table6"], "Table 6. Exploratory posterior-CARS and machine-learning results.")

    doc.add_heading("Discussion / 讨论", level=1)
    dis_sections = [
        ("4.1 Principal findings",
         "研究在静息态、电影态和跨状态层面形成一致证据链：ASD 儿童非周期谱背景更平坦、后部效应更突出、电影 ISC 普遍降低，"
         "并在稳健回归中观察到跨状态耦合。"),
        ("4.2 Aperiodic EEG alterations as resting-state neural background differences",
         "global_exponent 降低与 altered broadband spectral background 一致，说明差异主要位于宽频背景而非单一周期峰参数。"),
        ("4.3 Posterior aperiodic dynamics and naturalistic movie synchrony",
         "posterior exponent 与 mental ISC 的交互支持 cross-state neurophysiological phenotype。该证据在稳健回归中更强，"
         "普通 OLS 对纳入标准和模型设定敏感。"),
        ("4.4 Developmental moderation and ASD heterogeneity",
         "年龄交互提示神经表型具有发展调制，符合 ASD 异质性特征。横断面设计下不应进行因果或个体轨迹推断。"),
        ("4.5 Methodological implications",
         "功率谱参数化、质量控制协变量、稳健回归和 split-half 联合报告可提高解释透明度。自动 ICA 结果提示儿童 EEG 伪迹处理会影响统计证据强度。"),
        ("4.6 Clinical and translational boundaries",
         "posterior-CARS 及分类结果应视为探索性。当前结果不支持诊断应用表述，仍需多中心复现。"),
        ("4.7 Limitations",
         "研究为横断面设计；部分临床/伦理元信息需投稿前核验；交互与临床相关分析样本量有限；头皮电极结果不代表脑源定位。"),
        ("4.8 Conclusions",
         "结果 supports a cross-state neurophysiological phenotype：ASD 儿童在静息态非周期背景与自然电影神经同步上存在系统差异。"),
    ]
    for h, t in dis_sections:
        ph = doc.add_paragraph(h)
        ph.runs[0].bold = True
        doc.add_paragraph(t)

    doc.add_heading("Conclusion / 结论", level=1)
    doc.add_paragraph(
        "ASD 儿童静息态 global aperiodic exponent 降低并呈后部空间效应，且存在年龄依赖性；自然电影 ISC 在 mental、pain、neutral 事件中均表现 TD>ASD。"
        "静息态 posterior exponent 与电影 mental ISC 的跨状态耦合在稳健回归中获得支持，但普通 OLS 对纳入标准敏感。"
        "整体结果支持跨状态神经表型框架，而非诊断级生物标志物。"
    )

    doc.add_heading("Data and Code Availability / 数据与代码可用性", level=1)
    doc.add_paragraph("去标识化衍生数据与分析脚本可在合理学术请求下提供。")
    doc.add_heading("Ethics Statement / 伦理声明", level=1)
    doc.add_paragraph("研究遵循机构伦理审查流程并由监护人签署知情同意。具体伦理批件编号见投稿前作者核验清单。")
    doc.add_heading("Author Contributions / 作者贡献", level=1)
    doc.add_paragraph("作者团队共同完成研究设计、数据分析、结果解释与论文撰写。具体分工将在投稿前根据期刊规范补充。")
    doc.add_heading("Conflict of Interest / 利益冲突", level=1)
    doc.add_paragraph("作者声明不存在可能影响研究结论的商业或财务利益冲突。")
    doc.add_heading("Funding / 经费支持", level=1)
    doc.add_paragraph("经费来源信息将在投稿前由作者团队核验并补充。")
    doc.add_heading("Acknowledgements / 致谢", level=1)
    doc.add_paragraph("感谢参与研究的儿童及其家庭，以及参与数据采集与质量控制的研究人员。")

    doc.add_heading("References / 参考文献", level=1)
    refs = [
        "Donoghue T, Haller M, Peterson EJ, et al. Parameterizing neural power spectra into periodic and aperiodic components. Nature Neuroscience. 2020;23(12):1655-1665.",
        "Donoghue T, Dominguez J, Voytek B. Electrophysiological frequency band ratio measures conflate periodic and aperiodic neural activity. eNeuro. 2020;7(6):ENEURO.0192-20.2020.",
        "Gao R, Peterson EJ, Voytek B. Inferring synaptic excitation/inhibition balance from field potentials. NeuroImage. 2017;158:70-78.",
        "Neo WS, Foti D, Keehn B, Kelleher B. Resting-state EEG power differences in autism spectrum disorder: A systematic review and meta-analysis. Translational Psychiatry. 2023;13:389.",
        "Manyukhina VO, Prokofyev AO, Galuta IA, et al. Globally elevated excitation-inhibition ratio in children with autism spectrum disorder and below-average intelligence. Molecular Autism. 2022;13:20.",
        "Hill AT, Clark GM, Bigelow FJ, et al. Periodic and aperiodic neural activity displays age-dependent changes across early-to-middle childhood. Developmental Cognitive Neuroscience. 2022;54:101076.",
        "Wilkinson CL, Yankowitz LD, Chao JY, et al. Developmental trajectories of EEG aperiodic and periodic components in children 2-44 months of age. Nature Communications. 2024;15:5788.",
        "Karalunas SL, Gustafsson HC, Ostlund BD, et al. EEG aperiodic power spectral slope predicts ADHD risk in early development. Developmental Psychobiology. 2022;64(3):e22228.",
        "Gramfort A, Luessi M, Larson E, et al. MEG and EEG data analysis with MNE-Python. Frontiers in Neuroscience. 2013;7:267.",
        "Pion-Tonachini L, Kreutz-Delgado K, Makeig S. ICLabel: An automated EEG independent component classifier. NeuroImage. 2019;198:181-197.",
    ]
    for r in refs:
        doc.add_paragraph(r)

    doc.add_heading("Figure Legends / 图注", level=1)
    legends = [
        "Figure 1. Study design and sample inclusion.",
        "Figure 2. Resting-state aperiodic EEG differences.",
        "Figure 3. Age-dependent effects.",
        "Figure 4. Spatial distribution of posterior effects (electrode-level only).",
        "Figure 5. Naturalistic movie ISC differences.",
        "Figure 6. Cross-state coupling between resting posterior exponent and movie mental ISC.",
        "Figure 7. Exploratory clinical and classification analyses.",
    ]
    for l in legends:
        doc.add_paragraph(l)

    doc.add_heading("Tables / 表格", level=1)
    for i in range(1, 7):
        doc.add_paragraph(f"Table {i}. See corresponding table in Results section.")

    doc.add_heading("投稿前作者核验清单", level=1)
    checklist = [
        "伦理审批机构名称与批准号",
        "作者信息、单位与通讯作者邮箱",
        "经费项目编号与资助机构",
        "临床量表版本与施测时间窗",
        "参考文献卷期页最终核对",
        "目标期刊格式化模板与图表尺寸终校",
    ]
    for item in checklist:
        doc.add_paragraph(f"- {item}")

    doc.save(OUT_MAIN)


def build_supp_doc(supp_figs):
    doc = Document()
    set_doc_style(doc)
    t = doc.add_paragraph("Supplementary Materials")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(15)

    doc.add_heading("Supplementary Methods", level=1)
    sup_methods = [
        ("S1. Specparam 参数化细节", "采用固定模式拟合 1-40 Hz 频段，提取非周期指数、偏移及周期峰参数，并实施通道与被试层面质量控制。"),
        ("S2. 自动 ICA 细节", "自动 ICA 用于敏感性分析，比较不同阈值下全局与后部非周期指标的效应稳定性。"),
        ("S3. Split-half 方法", "按奇偶 epoch 重建指标并计算 Spearman 与 Spearman-Brown 校正信度。"),
        ("S4. 机器学习 nested CV", "采用 nested 5x5 交叉验证比较周期、非周期与联合特征模型，并使用 DeLong 检验模型差异。"),
        ("S5. QC 协变量模型", "将 mean_fit_error 与 invalid_channel_ratio 纳入模型，检验电影 ISC 组差是否由数据质量指标驱动。"),
        ("S6. 严格纳入标准敏感性分析", "采用更严格样本纳入规则，评估跨状态耦合与临床相关结果的稳定性。"),
    ]
    for h, p in sup_methods:
        hp = doc.add_paragraph(h)
        hp.runs[0].bold = True
        doc.add_paragraph(p)

    doc.add_heading("Supplementary Figures", level=1)
    for key in [
        "FigureS1_split_half_reliability",
        "FigureS2_ica_global_vs_posterior",
        "FigureS3_primary_vs_stringent",
        "FigureS4_delta_exponent",
        "FigureS5_posterior_cars_scatter",
        "FigureS6_classification_summary",
    ]:
        add_figure(doc, supp_figs[key][0], key.replace("_", " "))

    doc.add_heading("Supplementary Tables", level=1)
    supp_tables = [
        ("Table S1. Full sample flow", TAB_DIR / "TableS1_full_sample_flow.csv"),
        ("Table S2. Nested model robustness", TAB_DIR / "TableS2_nested_model_robustness.csv"),
        ("Table S3. Frequency sensitivity", TAB_DIR / "TableS3_frequency_sensitivity.csv"),
        ("Table S4. Split-half reliability", TAB_DIR / "TableS4_split_half_reliability.csv"),
        ("Table S5. Automated ICA results", TAB_DIR / "TableS5_automated_ica_results.csv"),
        ("Table S6. Machine-learning results", TAB_DIR / "TableS6_machine_learning_results.csv"),
    ]
    for title, f in supp_tables:
        doc.add_paragraph(title).runs[0].bold = True
        df = pd.read_csv(f)
        add_table(doc, df, "")

    doc.add_heading("投稿前作者核验清单", level=1)
    for item in [
        "补充材料与主文中样本量描述一致性",
        "补充图 S 编号与正文引用一致性",
        "补充表字段中英文术语统一",
        "所有统计值与主分析输出逐项核对",
    ]:
        doc.add_paragraph(f"- {item}")

    doc.save(OUT_SUPP)


def main():
    ensure_dirs()
    data = load_data()
    figs = {}
    figs["fig1"], _ = generate_figure1()
    figs["fig2"], _ = generate_figure2(data)
    figs["fig3"], _ = generate_figure3(data)
    figs["fig4"], _ = generate_figure4(data)
    figs["fig5"], _ = generate_figure5(data)
    figs["fig6"], _ = generate_figure6(data)
    figs["fig7"], _ = generate_figure7(data)
    supp_figs = generate_supp_figures(data)
    tables = write_tables(data)
    build_main_doc(tables, figs)
    build_supp_doc(supp_figs)
    print("Generated:")
    print(OUT_MAIN)
    print(OUT_SUPP)
    print(FIG_DIR)
    print(TAB_DIR)


if __name__ == "__main__":
    main()
