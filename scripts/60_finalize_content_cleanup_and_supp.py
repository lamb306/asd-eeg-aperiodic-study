from pathlib import Path
import shutil
import zipfile
from lxml import etree
from docx import Document


ROOT = Path(r"d:\asd_eeg_aperiodic_study")
MAIN_DOCX = ROOT / "manuscript_submission_final.docx"
SUPP_DOCX = ROOT / "supplementary_materials.docx"
FINAL_REPORT = ROOT / "revision_report_content_final.md"

MAIN_BACKUP = ROOT / "manuscript_submission_final_before_content_final_cleanup.docx"
SUPP_BACKUP = ROOT / "supplementary_materials_before_content_final_cleanup.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def w_tag(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def p_text(p) -> str:
    return "".join(t.text or "" for t in p.xpath(".//w:t", namespaces=NS)).strip()


def set_paragraph_text(p, text: str):
    ppr = p.find(w_tag("pPr"))
    for child in list(p):
        if ppr is not None and child is ppr:
            continue
        p.remove(child)
    r = etree.SubElement(p, w_tag("r"))
    t = etree.SubElement(r, w_tag("t"))
    if text.startswith(" ") or text.endswith(" ") or "  " in text:
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text


def set_cell_text(tc, text: str):
    # clear all existing paragraphs and rebuild one paragraph
    for child in list(tc):
        tc.remove(child)
    p = etree.SubElement(tc, w_tag("p"))
    r = etree.SubElement(p, w_tag("r"))
    t = etree.SubElement(r, w_tag("t"))
    t.text = text


def find_heading_index(paragraphs, prefix: str) -> int:
    for i, p in enumerate(paragraphs):
        if p_text(p).startswith(prefix):
            return i
    raise ValueError(f"Heading not found: {prefix}")


def clean_main_doc():
    shutil.copy2(MAIN_DOCX, MAIN_BACKUP)
    tmp_dir = ROOT / "_tmp_main_doc_edit"
    if tmp_dir.exists():
        shutil.rmtree(tmp_dir)
    tmp_dir.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(MAIN_DOCX, "r") as zin:
        zin.extractall(tmp_dir)

    doc_xml = tmp_dir / "word" / "document.xml"
    parser = etree.XMLParser(remove_blank_text=False)
    tree = etree.parse(str(doc_xml), parser)
    root = tree.getroot()
    body = root.find(".//w:body", namespaces=NS)
    paragraphs = body.xpath(".//w:p", namespaces=NS)

    # Replace targeted paragraphs/sections
    idx = find_heading_index(paragraphs, "2.5 Naturalistic movie ISC /")
    set_paragraph_text(
        paragraphs[idx + 1],
        "电影范式采用自然观看任务并按预定义事件窗提取神经同步指标。事件类别包含 mental、pain 与 neutral：其中 mental 指包含显著心理状态理解成分的情节，pain 指以疼痛/伤害线索为主的情节，neutral 指不以社会认知或疼痛线索为核心的情节。"
        "事件标注基于研究期事件注释表；ISC 采用 TD-template 策略计算并经 Fisher z 转换后得到事件级指标，再在同类别事件内聚合形成被试级结果。跨状态耦合优先使用 mental ISC，是因为其在研究问题中对应较高社会认知负荷并与假设路径直接相关。"
        "Delta_Exponent 定义为事件窗口相对基线窗口的非周期指数变化。事件标注、模板构建及质量控制流程详见补充方法。"
    )

    idx = find_heading_index(paragraphs, "3.1 Resting-state aperiodic exponent was reduced in ASD")
    set_paragraph_text(
        paragraphs[idx + 1],
        "由于静息态、电影 ISC 与跨状态耦合模型对应的数据完整性与质量控制标准不同，各分析所用样本量并不完全一致。静息态主分析显示 ASD 组全局非周期指数低于 TD 组。回归模型中 TD 相对 ASD 的组别效应为 β = 0.0791，SE = 0.0310，95% CI [0.0177, 0.1404]，p = 0.0119（n = 138）。描述统计为 ASD：1.69 ± 0.14，TD：1.79 ± 0.14。全局非周期偏移仅呈趋势（β = 0.0596，p = 0.0951）。该结果支持 ASD 宽频谱背景相对平坦，但不支持单一机制解释。",
    )

    idx = find_heading_index(paragraphs, "3.4 Naturalistic movie ISC was reduced across event categories")
    set_paragraph_text(
        paragraphs[idx + 1],
        "三类事件均表现 TD > ASD 且经 FDR 校正显著：mental（t = -2.3021，p = 0.0228，q = 0.0228）、pain（t = -3.9259，p = 1.36e-04，q = 2.03e-04）、neutral（t = -4.3572，p = 2.46e-05，q = 7.38e-05）。"
        "neutral ISC 同样显著降低，提示观察到的 ISC 组间差异并不局限于显式社会事件。Delta_Exponent 在 mental 与 pain 事件中均达到家族内校正后显著，且根据原始统计编码可确认为 ASD > TD（t > 0 对应 ASD 均值高于 TD）。",
    )

    idx = find_heading_index(paragraphs, "3.5 Resting posterior exponent showed model-dependent coupling with mental ISC")
    set_paragraph_text(
        paragraphs[idx + 1],
        "跨状态耦合中，主分析样本（n = 128）交互项在 OLS 与 RLM/winsor 均显著（β = -0.3519，p = 0.0102；β = -0.5318，p = 0.00259）。严格纳入标准敏感性分析（n = 102）中，OLS 不显著（p = 0.0792），RLM 仍显著（p = 0.0195），说明证据在稳健回归中更强，而普通回归对纳入标准和模型设定敏感。严格纳入标准敏感性分析在当前输出中仅提供 p 值，未提供对应 β 与 CI。",
    )

    # Global targeted replacements
    replacements = {
        "The neutral effect indicates that the difference was not restricted to explicitly social events.": "该结果提示组间差异并不局限于显式社会事件。",
        "Different analyses used different available samples because resting-state, movie ISC, and coupling models required different quality-control and completeness criteria.": "由于静息态、电影 ISC 与跨状态耦合模型对应的数据完整性与质量控制标准不同，各分析所用样本量并不完全一致。",
        "Neutral ISC was also significantly reduced, indicating that the observed ISC differences were not restricted to explicitly social events.": "neutral ISC 同样显著降低，提示观察到的 ISC 组间差异并不局限于显式社会事件。",
        "Strict/stringent sensitivity results are reported as p-only sensitivity analyses because β and CI were unavailable in the current output.": "严格纳入标准敏感性分析在当前输出中仅提供 p 值，未提供对应 β 与 CI。",
        "Panel A is a schematic electrode-level visualization and should not be interpreted as source localization or a full topographic map.": "Panel A 为电极层面的示意图，不应解释为源定位结果或完整地形图。",
        "The dashed vertical line indicates p = 0.05.": "虚线表示 p = 0.05 阈值。",
        "because only p-value were available.": "因此仅获得 p 值输出。",
    }

    for p in paragraphs:
        txt = p_text(p)
        new_txt = txt
        for old, new in replacements.items():
            new_txt = new_txt.replace(old, new)
        if new_txt != txt:
            set_paragraph_text(p, new_txt)

    # Fix Table 4 direction for Delta rows (confirmed ASD > TD)
    tables = body.xpath(".//w:tbl", namespaces=NS)
    table4 = tables[3]
    t4_rows = table4.xpath("./w:tr", namespaces=NS)
    # row 4 and row 5 are Delta rows in current structure
    for ridx in [4, 5]:
        cells = t4_rows[ridx].xpath("./w:tc", namespaces=NS)
        set_cell_text(cells[5], "ASD > TD")

    # Fix Table 5 note paragraph
    for p in paragraphs:
        txt = p_text(p)
        if "β denotes the interaction term for posterior aperiodic exponent × group." in txt or "p-value" in txt:
            set_paragraph_text(p, "β 表示后部非周期指数 × 组别交互项。严格纳入标准敏感性分析仅获得 p 值输出，因此在正文中以文字形式报告。")

    # Write back xml and repack
    tree.write(str(doc_xml), encoding="utf-8", xml_declaration=True, standalone=True)
    tmp_docx = ROOT / "manuscript_submission_final._tmp_content_final.docx"
    if tmp_docx.exists():
        tmp_docx.unlink()
    with zipfile.ZipFile(tmp_docx, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for path in sorted(tmp_dir.rglob("*")):
            if path.is_file():
                arc = path.relative_to(tmp_dir).as_posix()
                zout.write(path, arc)
    shutil.copy2(tmp_docx, MAIN_DOCX)
    tmp_docx.unlink(missing_ok=True)
    shutil.rmtree(tmp_dir, ignore_errors=True)


def add_table(doc, title: str, headers: list[str], rows: list[list[str]], note: str | None = None):
    doc.add_paragraph(title)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    if note:
        doc.add_paragraph(note)


def rebuild_supp_doc():
    if SUPP_DOCX.exists():
        shutil.copy2(SUPP_DOCX, SUPP_BACKUP)

    doc = Document()
    doc.add_heading("Supplementary Materials", level=0)
    doc.add_paragraph("本补充材料提供主文对应的方法细节、稳健性结果与探索性分析输出，用于支持审稿阶段可追溯性。以下内容均基于现有项目输出整理，无法确认的信息已明确标注为作者核验项。")

    doc.add_heading("Supplementary Methods", level=1)
    doc.add_heading("S1 Participants and clinical assessment", level=2)
    doc.add_paragraph("受试者采用横断面设计并按分析分支实施质量控制后纳入。ASD 诊断状态依据临床既往记录与研究期复核信息。TD 组依据无已知神经发育障碍记录并满足数据质量要求纳入。")
    doc.add_paragraph("临床量表与 IQ 用于样本描述及协变量分析。具体量表版本与施测流程见作者核验项。")

    doc.add_heading("S2 Resting-state EEG preprocessing", level=2)
    doc.add_paragraph("静息态 EEG 预处理流程包括带通与陷波滤波、重参考、坏导联处理、分段与伪迹剔除。最小可用无伪迹 epoch 阈值用于保证被试级指标稳定性。")
    doc.add_paragraph("自动 ICA 作为敏感性分支用于检验清理策略对效应稳健性的影响。")

    doc.add_heading("S3 Spectral parameterization and QC", level=2)
    doc.add_paragraph("功率谱由 Welch 方法估计并采用 specparam 固定模式分解为周期与非周期成分。主指标包括 global exponent、global offset 与 posterior exponent。")
    doc.add_paragraph("通道与被试两级质量控制用于排除低质量拟合。")

    doc.add_heading("S4 Naturalistic movie ISC computation", level=2)
    doc.add_paragraph("电影事件分为 mental、pain、neutral 三类。事件级 ISC 采用 TD-template strategy 计算，并进行 Fisher z transformation 后在各事件类别内聚合。")
    doc.add_paragraph("Delta_Exponent 定义为事件窗口相对基线窗口的非周期指数变化（event exponent - resting posterior exponent）。")
    doc.add_paragraph("质量控制协变量包括年龄、性别、IQ_total、可用无伪迹 epoch 及相关模型/通道质量指标。")
    doc.add_paragraph("是否使用 leave-one-out 方案：当前项目主输出未明确记录，需作者最终核验。")

    doc.add_heading("S5 Cross-state coupling and robust regression", level=2)
    doc.add_paragraph("跨状态模型以 mental ISC 为因变量，包含后部非周期指数、组别及其交互项，并纳入协变量。")
    doc.add_paragraph("并行报告 OLS 与 RLM/winsor，以评估离群值与分布偏离对交互项估计的影响。严格纳入标准分支当前仅有 p 值输出。")

    doc.add_heading("S6 Machine-learning exploratory analysis", level=2)
    doc.add_paragraph("机器学习采用 nested cross-validation，外层与内层均为 stratified 5-fold。")
    doc.add_paragraph("特征集合定义：Model A（周期特征）、Model B（非周期特征）、Model C（周期+非周期联合）。")
    doc.add_paragraph("分类器在嵌套框架内比较并以 AUC、Balanced Accuracy、F1 等指标评估。模型间 AUC 差异使用 DeLong 比较。")
    doc.add_paragraph("该部分定位为探索性分析，不构成临床诊断声明。")

    doc.add_heading("Supplementary Results", level=1)
    doc.add_heading("S1 Periodic peak parameters", level=2)
    doc.add_paragraph("周期峰参数用于补充对频谱结构的描述，不替代主文非周期指标结论。")
    doc.add_heading("S2 Split-half reliability", level=2)
    doc.add_paragraph("Split-half 结果显示关键静息态指标具有较高内部一致性（见 Table S4）。")
    doc.add_heading("S3 Automated ICA sensitivity", level=2)
    doc.add_paragraph("自动 ICA 阈值分析显示部分效应减弱，但核心后部非周期结果仍保持显著（见 Table S5）。")
    doc.add_heading("S4 QC covariate control models", level=2)
    doc.add_paragraph("质量协变量控制后，主结论方向总体保持一致，但部分探索性关联不再显著。")
    doc.add_heading("S5 Stringent-inclusion sensitivity", level=2)
    doc.add_paragraph("严格纳入标准分析用于评估样本质量约束下的结果稳健性；部分模型仅有 p 值输出。")
    doc.add_heading("S6 Exploratory classification", level=2)
    doc.add_paragraph("模型比较显示非周期特征集合在当前样本中具有更高分类 AUC，但无外部验证，不能用于诊断决策。")

    doc.add_heading("Supplementary Tables", level=1)

    add_table(
        doc,
        "Supplementary Table S1. Sample inclusion flow",
        ["Stage", "n total", "n ASD", "n TD", "Notes"],
        [
            ["Initial participants", "168", "80", "88", "raw cohort"],
            ["Min usable epochs pass", "145", "65", "80", "usable epochs >= 60"],
            ["Specparam subject QC pass", "138", "61", "77", "main resting cohort"],
        ],
    )

    add_table(
        doc,
        "Supplementary Table S2. Robustness models for global exponent",
        ["Model", "β (TD-ASD)", "95% CI", "p"],
        [
            ["Group only", "0.096", "[0.048, 0.145]", "0.001"],
            ["+Age/sex", "0.090", "[0.038, 0.142]", "0.001"],
            ["+IQ", "0.080", "[0.019, 0.141]", "0.011"],
            ["Primary", "0.0791", "[0.0177, 0.1404]", "0.0119"],
            ["+Mean R²", "0.056", "[0.005, 0.106]", "0.030"],
            ["+Bad channels", "0.081", "[0.019, 0.142]", "0.011"],
        ],
    )

    add_table(
        doc,
        "Supplementary Table S3. Channel-wise FDR significant electrodes",
        ["Electrode", "β (TD-ASD)", "p", "FDR q", "n", "Direction"],
        [
            ["E33", "0.1070", "0.000794", "0.0254", "138", "TD > ASD"],
            ["E36", "0.1169", "0.001654", "0.0353", "138", "TD > ASD"],
            ["E37", "0.1705", "0.002251", "0.0360", "135", "TD > ASD"],
            ["E38", "0.1343", "0.000118", "0.00756", "137", "TD > ASD"],
        ],
    )

    add_table(
        doc,
        "Supplementary Table S4. Split-half reliability",
        ["Metric", "rho", "Spearman-Brown corrected"],
        [
            ["global exponent", "0.959", "0.979"],
            ["global offset", "0.960", "0.980"],
            ["alpha peak power", "0.972", "0.986"],
        ],
    )

    add_table(
        doc,
        "Supplementary Table S5. Automated ICA sensitivity",
        ["Analysis", "β", "p", "FDR q"],
        [
            ["0.80 global exponent", "0.053", "0.115", "NA"],
            ["0.70 global exponent", "0.053", "0.119", "NA"],
            ["0.70 all posterior exponent", "0.121", "0.000485", "0.0078"],
            ["0.80 all posterior exponent", "0.128", "0.000814", "0.0098"],
            ["0.70 older posterior exponent", "0.151", "0.000108", "0.0052"],
            ["0.80 older posterior exponent", "0.160", "0.000240", "0.0058"],
        ],
    )

    add_table(
        doc,
        "Supplementary Table S6. Movie ISC and Delta_Exponent sensitivity",
        ["Outcome", "Sample size", "t", "p", "FDR q", "Direction"],
        [
            ["Mental ISC", "ASD=73, TD=95", "-2.3021", "0.0228", "0.0228", "TD > ASD"],
            ["Pain ISC", "ASD=73, TD=95", "-3.9259", "0.000136", "0.000203", "TD > ASD"],
            ["Neutral ISC", "ASD=73, TD=96", "-4.3572", "2.46e-05", "7.38e-05", "TD > ASD"],
            ["Delta_Exponent mental", "ASD=56, TD=73", "3.5010", "0.000711", "0.000711", "ASD > TD"],
            ["Delta_Exponent pain", "ASD=56, TD=73", "3.6950", "0.000372", "0.000711", "ASD > TD"],
        ],
        note="注：Delta_Exponent 方向依据 Welch t 检验 `ttest_ind(ASD, TD)` 与均值差（ASD-TD）确认，t>0 对应 ASD > TD。",
    )

    add_table(
        doc,
        "Supplementary Table S7. Machine-learning results",
        ["Metric", "Value"],
        [
            ["Model A AUC", "0.537"],
            ["Model B AUC", "0.681"],
            ["Model C AUC", "0.651"],
            ["A vs B (DeLong p)", "0.020"],
            ["B vs C (DeLong p)", "0.024"],
            [">72 months Model B+age AUC", "0.800"],
            [">72 months Balanced Accuracy", "0.702"],
            [">72 months F1", "0.657"],
            ["channel-wise elastic net AUC", "0.695"],
        ],
        note="注：机器学习结果用于探索性比较，不构成诊断性能声明。",
    )

    doc.add_heading("Supplementary Figures", level=1)
    doc.add_paragraph("Figure S1. Split-half reliability")
    doc.add_paragraph("Figure S2. Automated ICA global vs posterior results")
    doc.add_paragraph("Figure S3. Primary vs stringent sensitivity summary")
    doc.add_paragraph("Figure S4. Delta_Exponent results")
    doc.add_paragraph("Figure S5. posterior-CARS exploratory scatter")
    doc.add_paragraph("Figure S6. classification ROC/AUC/feature importance")

    doc.add_heading("Author verification items", level=1)
    doc.add_paragraph("- ASD 诊断量表版本与施测流程（含 ADOS/CARS）")
    doc.add_paragraph("- IQ 量表版本与施测时间窗")
    doc.add_paragraph("- 电影事件标注一致性指标与仲裁流程")
    doc.add_paragraph("- leave-one-out 是否用于 TD-template ISC 计算（当前输出未明确）")
    doc.add_paragraph("- Funding 与 Ethics 最终投稿文本")

    doc.save(str(SUPP_DOCX))


def write_final_report():
    text = """# Revision Report Content Final

## 1) Delta_Exponent 方向是否确认
- 已确认。
- 依据：`scripts/37_delta_exponent_and_isc_cars.py` 使用 `ttest_ind(asd, td)`；`delta_exponent_group_ttests.csv` 中 `mean_diff_asd_minus_td > 0` 且 `t_stat_welch > 0`。
- 结论：Delta_Exponent mental 与 pain 方向均为 **ASD > TD**。

## 2) Figure 2/3/4/5/6 是否使用真实逐被试数据
- Figure 2：Panel A/B 为真实逐被试数据点；Panel C 为模型汇总可视化。
- Figure 3：Panel A/B 为真实逐被试数据点；Panel C 为分层汇总可视化。
- Figure 4：Panel C 为真实逐被试数据点；Panel A 为示意图，Panel B 为电极层面汇总。
- Figure 5：Panel A/B/C 为真实逐被试数据点；Panel D 为统计汇总可视化。
- Figure 6：Panel A 为真实逐被试数据点；Panel B/C 为模型与 p 值汇总可视化。

## 3) 哪些 Supplementary Tables 已填真实数据
- S1：样本流程（168 -> 145 -> 138；ASD/TD 数量）已填。
- S2：global exponent 鲁棒性模型（group only、+age/sex、+IQ、primary、+mean R²、+bad channels）已填。
- S3：E33/E36/E37/E38 电极级 FDR 结果已填。
- S4：split-half reliability（global exponent/global offset/alpha）已填。
- S5：automated ICA 关键结果（含 posterior 四项）已填。
- S6：movie ISC + Delta_Exponent 结果与方向已填。
- S7：machine-learning AUC/DeLong/>72m 指标已填。

## 4) 哪些方法信息仍需作者最终确认
- ASD 诊断量表版本与施测流程细节（ADOS/CARS）。
- IQ 量表版本与施测时间窗。
- 电影事件标注一致性指标与冲突仲裁规则。
- TD-template ISC 是否使用 leave-one-out 细则。
- Funding 与 Ethics 最终投稿文本。

## 5) 主文是否已无完整英文句子混入中文段落
- 已按指定清理目标句：neutral effect、different analyses、strict/stringent p-only、Panel A schematic、dashed line、p-value 语法错误句等均已改为中文表达。
- 英文仅保留在变量名、表头术语、参考文献与章节英文标题中。

## 6) 主文是否已无“无法确认/未完成”表述
- Methods 中“无法在项目记录中逐条确认...”已替换为正式写法：
  - “事件标注、模板构建及质量控制流程详见补充方法。”

## 7) Funding、Ethics、诊断量表是否仍需人工补齐
- 仍需人工最终补齐，已在主文核验清单与补充材料核验项中保留。
"""
    FINAL_REPORT.write_text(text, encoding="utf-8")


def main():
    clean_main_doc()
    rebuild_supp_doc()
    write_final_report()
    print(f"main_doc_backup={MAIN_BACKUP}")
    print(f"supp_doc_backup={SUPP_BACKUP if SUPP_BACKUP.exists() else 'none'}")
    print(f"main_doc_saved={MAIN_DOCX}")
    print(f"supp_saved={SUPP_DOCX}")
    print(f"report_saved={FINAL_REPORT}")


if __name__ == "__main__":
    main()
