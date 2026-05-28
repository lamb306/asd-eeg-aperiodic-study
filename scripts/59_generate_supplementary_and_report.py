from pathlib import Path
import shutil
from docx import Document


ROOT = Path(r"d:\asd_eeg_aperiodic_study")
SUPP = ROOT / "supplementary_materials.docx"
SUPP_BAK = ROOT / "supplementary_materials_before_content_strengthen.docx"
REPORT = ROOT / "revision_report_content.md"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it)
        p.style = "List Bullet"


def build_supplementary():
    if SUPP.exists():
        shutil.copy2(SUPP, SUPP_BAK)

    doc = Document()
    add_heading(doc, "Supplementary Materials", level=0)
    add_para(doc, "This document provides extended methodological details, sensitivity results, and exploratory outputs that support the main manuscript without changing the prespecified primary conclusions.")
    add_para(doc, "Where source records are incomplete or version metadata are unavailable, we explicitly flag the missing items instead of imputing or fabricating details.")

    add_heading(doc, "Supplementary Methods", level=1)
    add_heading(doc, "S1 Participants and clinical assessment", level=2)
    add_para(doc, "Participants were analyzed under a cross-sectional design with task-specific inclusion after quality control. ASD diagnostic status was taken from available clinical records and study-period verification notes using conservative wording in the main text.")
    add_para(doc, "TD inclusion required no known neurodevelopmental diagnosis in available records and adequate data quality for the corresponding analysis branch.")
    add_para(doc, "Exclusions were applied for insufficient artifact-free EEG data, failed spectral-parameterization QC, or missing key analysis variables.")
    add_para(doc, "IQ and clinical scales were used as descriptive and sensitivity covariates. Specific instrument names/versions are pending author verification.")

    add_heading(doc, "S2 Resting-state EEG preprocessing", level=2)
    add_para(doc, "The preprocessing stream followed an MNE-Python based workflow including band-pass filtering, notch filtering, downsampling, bad-channel handling, rereferencing, and fixed-length epoching.")
    add_para(doc, "A minimum number of artifact-free epochs per subject was enforced before resting-state metrics were estimated.")
    add_para(doc, "Automated ICA branches were evaluated as sensitivity analyses rather than primary preprocessing.")

    add_heading(doc, "S3 Spectral parameterization and QC", level=2)
    add_para(doc, "Power spectra were estimated and decomposed into aperiodic and periodic components using specparam fixed mode for primary analyses.")
    add_para(doc, "Primary metrics included global exponent, global offset, and posterior exponent.")
    add_para(doc, "Channel-level and subject-level QC thresholds were applied before model fitting. Knee-mode and alternative cleaning settings were treated as sensitivity checks.")

    add_heading(doc, "S4 Naturalistic movie ISC computation", level=2)
    add_para(doc, "Movie events were grouped into mental, pain, and neutral categories from the project event-annotation table.")
    add_para(doc, "ISC was computed using a TD-template strategy, Fisher z transformed, and aggregated to event-level and category-level subject summaries.")
    add_para(doc, "Delta_Exponent was defined as event-window minus baseline-window aperiodic exponent.")
    add_para(doc, "Annotation workflow details (annotator count, agreement metrics, conflict resolution) require final author confirmation.")

    add_heading(doc, "S5 Cross-state coupling and robust regression", level=2)
    add_para(doc, "Cross-state models linked posterior resting exponent to movie ISC with group interaction terms and prespecified covariates.")
    add_para(doc, "OLS and robust regression (RLM/winsor) were run in parallel to evaluate sensitivity to outliers and distributional violations.")
    add_para(doc, "For stringent-inclusion sensitivity outputs, only p-values were available in the locked output tables; beta and CI are therefore not reported for those rows.")

    add_heading(doc, "S6 Machine-learning exploratory analysis", level=2)
    add_para(doc, "Machine-learning analyses were exploratory and used nested cross-validation for internal model comparison.")
    add_para(doc, "These analyses were not designed or validated as clinical diagnostic tools.")

    add_heading(doc, "Supplementary Results", level=1)
    add_heading(doc, "S1 Periodic peak parameters", level=2)
    add_para(doc, "Periodic peak parameters were retained as secondary context; they did not replace or overturn the main aperiodic findings.")
    add_heading(doc, "S2 Split-half reliability", level=2)
    add_para(doc, "Split-half reliability analyses were used to assess internal stability of resting-state metrics.")
    add_heading(doc, "S3 Automated ICA sensitivity", level=2)
    add_para(doc, "Automated ICA thresholds (0.80 and 0.70) attenuated some effects but did not reverse the direction of primary group contrasts in available outputs.")
    add_heading(doc, "S4 QC covariate control models", level=2)
    add_para(doc, "Adding QC covariates tested whether findings were explained by data quality variation rather than group-related neural metrics.")
    add_heading(doc, "S5 Stringent-inclusion sensitivity", level=2)
    add_para(doc, "Stringent sensitivity models are reported as p-only where beta/CI were not available in current prespecified outputs.")
    add_heading(doc, "S6 Exploratory classification", level=2)
    add_para(doc, "Exploratory classifiers showed statistical separability patterns but remain non-diagnostic without external validation.")

    add_heading(doc, "Supplementary Tables", level=1)
    add_bullets(
        doc,
        [
            "S1 Sample inclusion flow",
            "S2 Robustness models for global exponent",
            "S3 Channel-wise FDR results",
            "S4 Split-half reliability",
            "S5 Automated ICA sensitivity",
            "S6 Movie ISC and Delta_Exponent sensitivity",
            "S7 Machine-learning results",
        ],
    )

    add_heading(doc, "Supplementary Figures", level=1)
    add_bullets(
        doc,
        [
            "S1 Split-half reliability",
            "S2 Automated ICA global vs posterior results",
            "S3 Primary vs stringent sensitivity summary",
            "S4 Delta_Exponent results",
            "S5 posterior-CARS exploratory scatter",
            "S6 classification ROC/AUC/feature importance",
        ],
    )

    add_heading(doc, "Items requiring author verification", level=1)
    add_bullets(
        doc,
        [
            "ASD diagnostic workflow details, including formal instrument versions and administration procedures.",
            "IQ instrument name/version and administration window.",
            "Movie annotation personnel, inter-rater reliability metric, and adjudication protocol.",
            "Exact contrast coding for Delta_Exponent t statistics (direction mapping).",
            "Final funding statement and grant identifiers.",
        ],
    )

    doc.save(str(SUPP))


def build_revision_report():
    text = """# Revision Report (Content Strengthening)

## Output files
- manuscript_submission_final.docx
- supplementary_materials.docx
- revision_report_content.md

## Methods information provenance
- **From existing manuscript/project records**
  - Sample flow numbers and task-specific sample sizes (resting/movie/coupling branches).
  - Core EEG preprocessing chain, specparam-based aperiodic metrics, and QC framing.
  - ISC event categories (mental/pain/neutral), Fisher-z transformation, and event aggregation framing.
  - Cross-state coupling model structure and robust-regression sensitivity rationale.
  - Prespecified p-only status for stringent sensitivity outputs.
- **Still requires author confirmation**
  - ASD diagnostic instrument versions and verification workflow details.
  - IQ instrument name/version and implementation details.
  - Movie event annotation reliability details (annotators/metrics/adjudication).
  - Exact Delta_Exponent contrast coding direction.
  - Final funding statement for formal submission.

## Figure data authenticity and visualization level
- **Real subject-level data used**
  - Figure 2A/2B (group distributions with subject points)
  - Figure 3A/3B
  - Figure 4C
  - Figure 5A/5B/5C
  - Figure 6A
- **Summary-level visualizations**
  - Figure 1 (flow/schematic + availability bar chart)
  - Figure 2C (model summary forest-style panel)
  - Figure 3C (strata summary)
  - Figure 4A (schematic) and 4B (electrode-level forest summary)
  - Figure 5D (test-statistic summary)
  - Figure 6B/6C (interaction and p-summary panels)

## Delta_Exponent direction
- Current main text reports **group difference present** for Delta_Exponent rows.
- Exact direction coding (`t > 0` mapping to `TD > ASD` or `ASD > TD`) is **not confirmed from locked source metadata** and remains in author-verification list.

## Inference vs original record
- **Directly record-backed**: all numeric primary results quoted in main Results sections.
- **Conservative inferential wording added**: rationale paragraphs explaining event-category interpretation, robust-regression purpose, and translational boundaries.
- **No fabricated missing metadata**: missing instrument/version/coding details were explicitly moved to verification checklists instead of being invented.
"""
    REPORT.write_text(text, encoding="utf-8")


def main():
    build_supplementary()
    build_revision_report()
    print(f"supplementary_saved={SUPP}")
    print(f"supplementary_backup={SUPP_BAK if SUPP_BAK.exists() else 'none'}")
    print(f"report_saved={REPORT}")


if __name__ == "__main__":
    main()
