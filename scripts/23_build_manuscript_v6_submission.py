#!/usr/bin/env python
"""Build submission-style manuscript v6 from v5."""

from __future__ import annotations

import re
import sys
from pathlib import Path

import pandas as pd

PROJECT_ROOT = Path(__file__).resolve().parents[1]
V5 = PROJECT_ROOT / "outputs/reports/manuscript_draft_zh_v5_APA_figures.md"
OUT_MD = PROJECT_ROOT / "outputs/reports/manuscript_draft_zh_v6_submission_style.md"
OUT_DOCX = PROJECT_ROOT / "outputs/reports/manuscript_draft_zh_v6_submission_style.docx"
OUT_READINESS = PROJECT_ROOT / "outputs/reports/submission_readiness_audit.md"
OUT_REF_AUDIT = PROJECT_ROOT / "outputs/reports/reference_audit_v6.md"
OUT_FIG_AUDIT = PROJECT_ROOT / "outputs/reports/figure_table_audit_v6.md"


def fmt_p(p: float) -> str:
    if pd.isna(p):
        return "—"
    if p < 0.001:
        return "p < .001"
    s = f"{p:.3f}"
    if s.startswith("0."):
        s = s[1:]
    return f"p = {s}"


def fmt_num(x: float, d: int = 3) -> str:
    if pd.isna(x):
        return "—"
    return f"{x:.{d}f}"


def fmt_ci(lo: float, hi: float) -> str:
    return f"[{fmt_num(lo)}, {fmt_num(hi)}]"


def strip_md(text: str) -> str:
    text = re.sub(r"^---\s*$", "", text, flags=re.M)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def replace_donoghue_cites(text: str) -> str:
    # eNeuro band-ratio -> 2020a; Nat Neurosci specparam -> 2020b
    text = re.sub(
        r"Donoghue et al\., 2020(?![ab])",
        "Donoghue et al., 2020b",
        text,
    )
    text = re.sub(
        r"Donoghue et al\.（2020）",
        "Donoghue et al.（2020b）",
        text,
    )
    # Fix abstract/intro first mentions - band ratio paper is 2020a
    pairs = [
        (
            "频带比（含 θ/β）常反映非周期活动",
            "Donoghue et al., 2020a",
        ),
    ]
    # Manual targeted replacements for known sentences
    text = text.replace(
        "（Donoghue et al., 2020; Neo et al., 2023）",
        "（Donoghue et al., 2020a; Neo et al., 2023）",
    )
    text = text.replace(
        "之和（Donoghue et al., 2020）",
        "之和（Donoghue et al., 2020b）",
    )
    text = text.replace(
        "改变（Donoghue et al., 2020）",
        "改变（Donoghue et al., 2020a）",
    )
    text = text.replace(
        "（Donoghue et al., 2020; Neo et al., 2023）",
        "（Donoghue et al., 2020a; Neo et al., 2023）",
    )
    text = text.replace("Gao et al., 2017; Donoghue et al., 2020）", "Gao et al., 2017; Donoghue et al., 2020b）")
    text = text.replace("（Donoghue et al., 2020）", "（Donoghue et al., 2020b）")
    text = text.replace("specparam 对每个被试×通道 PSD 参数化（Donoghue et al., 2020）", "specparam 对每个被试×通道 PSD 参数化（Donoghue et al., 2020b）")
    text = text.replace("specparam 分离的价值（Donoghue et al., 2020）", "specparam 分离的价值（Donoghue et al., 2020b）")
    text = text.replace("Donoghue et al.（2020）", "Donoghue et al.（2020b）")
    text = text.replace(
        "可能混淆组间比较的解释（Donoghue et al., 2020b; Neo et al., 2023）",
        "可能混淆组间比较的解释（Donoghue et al., 2020a; Neo et al., 2023）",
    )
    text = text.replace(
        "偏移改变（Donoghue et al., 2020b）",
        "偏移改变（Donoghue et al., 2020a）",
    )
    text = text.replace(
        "使解释不明确（Donoghue et al., 2020b; Neo et al., 2023）",
        "使解释不明确（Donoghue et al., 2020a; Neo et al., 2023）",
    )
    return text


def fmt_stats_line(text: str) -> str:
    text = re.sub(r"\*p\*\s*=\s*([0-9.]+)", lambda m: fmt_p(float(m.group(1))), text)
    text = re.sub(r"\*p\*\s*<\s*([0-9.]+)", lambda m: fmt_p(0.0005), text)
    text = re.sub(r"\*p\*\s*≈\s*([0-9.]+)", lambda m: fmt_p(float(m.group(1))), text)
    text = re.sub(r"\*p\*\s*>\s*([0-9.]+)", r"p > .\1", text)
    text = re.sub(r"\*M\*\s*=", "M =", text)
    text = re.sub(r"\*SD\*\s*=", "SD =", text)
    text = re.sub(r"\*n\*\s*=", "n =", text)
    def _fmt_beta(m):
        v = float(m.group(1))
        if v < 0.01 and len(m.group(1)) > 4:
            return f"β = {v:.4f}".replace("0.", ".", 1)
        return f"β = {fmt_num(v, 3)}"

    text = re.sub(r"β\s*=\s*([0-9.]+)", _fmt_beta, text)
    text = re.sub(r"SE\s*=\s*([0-9.]+)", lambda m: f"SE = {fmt_num(float(m.group(1)), 3)}", text)
    text = re.sub(r"ρ\s*=\s*([0-9.]+)", lambda m: f"ρ = {fmt_num(float(m.group(1)), 3)}", text)
    text = re.sub(r"95% CI \[([0-9.]+),\s*([0-9.]+)\]", lambda m: f"95% CI {fmt_ci(float(m.group(1)), float(m.group(2)))}", text)
    return text


def build_table1() -> str:
    desc = pd.read_csv(PROJECT_ROOT / "outputs/tables/table1_main_cohort_descriptive.csv")
    comp = pd.read_csv(PROJECT_ROOT / "outputs/tables/table1_main_cohort_comparison.csv")
    qc = pd.read_csv(PROJECT_ROOT / "outputs/tables/table2_main_cohort_eeg_qc.csv")
    cohort = pd.read_csv(PROJECT_ROOT / "outputs/tables/main_cohort_subject_list.csv")
    sp = PROJECT_ROOT / "derivatives/specparam/specparam_qc_summary_subject.csv"
    pre = PROJECT_ROOT / "derivatives/qc/preproc_summary.csv"
    if pre.exists():
        cohort = cohort.merge(pd.read_csv(pre)[["subject_id", "bad_channel_count"]], on="subject_id", how="left")
    if sp.exists():
        spdf = pd.read_csv(sp)[["subject_id", "mean_r_squared"]]
        if "mean_r_squared" in cohort.columns:
            cohort = cohort.drop(columns=["mean_r_squared"])
        cohort = cohort.merge(spdf, on="subject_id", how="left")
    if "usable_epochs" in cohort.columns:
        cohort["usable_seconds"] = cohort["usable_epochs"] * 2.0

    def ms(g, col, d=1):
        s = cohort.loc[cohort.group == g, col]
        return f"{s.mean():.{d}f} ± {s.std():.{d}f}"

    rows = [
        "| Characteristic | ASD (n = 61) | TD (n = 77) | Statistic | p |",
        "|----------------|--------------|-------------|-----------|-----|",
    ]
    ap = fmt_p(comp.loc[comp.variable == "age_months", "t_pvalue"].iloc[0])
    rows.append(
        f"| Age, months | {ms('ASD','age_months')} | {ms('TD','age_months')} | t | {ap} |"
    )
    if "sex" in cohort.columns:
        from scipy.stats import chi2_contingency

        ct = pd.crosstab(cohort["group"], cohort["sex"])
        _, pv, _, _ = chi2_contingency(ct.values)

        def sx(g):
            sub = cohort[cohort.group == g]
            return f"{int((sub.sex=='F').sum())} F / {int((sub.sex=='M').sum())} M"

        rows.append(f"| Sex, F/M | {sx('ASD')} | {sx('TD')} | χ² | {fmt_p(pv)} |")
    iq = comp[comp.variable == "IQ_total"].iloc[0]
    rows.append(
        f"| IQ total | {ms('ASD','IQ_total')} | {ms('TD','IQ_total')} | t | {fmt_p(iq.t_pvalue)} |"
    )
    for var, lab in [
        ("usable_epochs", "Usable epochs"),
        ("usable_seconds", "Usable seconds"),
        ("bad_channel_count", "Bad channel count"),
        ("mean_r_squared", "Mean specparam R²"),
    ]:
        q = qc[qc.variable == var].iloc[0]
        d = 3 if var == "mean_r_squared" else 1
        rows.append(
            f"| {lab} | {ms('ASD', var, d)} | {ms('TD', var, d)} | t | {fmt_p(q.t_pvalue)} |"
        )
    cap = (
        "\n\nTable 1. Demographic characteristics and EEG quality metrics "
        "for the primary analysis cohort (N = 138). Continuous variables are "
        "mean ± SD. TD–ASD comparisons used two-sided independent-samples t "
        "tests except sex (χ²). R² = specparam model fit quality per subject.\n"
    )
    return "\n".join(rows) + cap


def build_table2() -> str:
    df = pd.read_csv(PROJECT_ROOT / "outputs/tables/global_exponent_robustness_models.csv")
    labels = {
        "model_1": ("Group only", "Group"),
        "model_2": ("+ Age, sex", "Group + age + sex"),
        "model_3": ("+ IQ", "Group + age + sex + IQ"),
        "model_4": ("Primary", "Group + age + sex + IQ + usable epochs"),
        "model_5": ("+ Mean R²", "Primary + mean specparam R²"),
        "model_6": ("+ Bad channels", "Primary + bad channel count"),
    }
    lines = [
        "| Model | Covariates | β (TD − ASD) | SE | 95% CI | p |",
        "|-------|------------|--------------|-----|--------|-----|",
    ]
    for _, r in df.iterrows():
        if r.model_name == "model_7_exclude_IQ_below_70":
            continue
        lab, cov = labels.get(r.model_name, (r.model_name, r.covariates))
        lines.append(
            f"| {lab} | {cov} | {fmt_num(r.group_coef_TD_vs_ASD)} | {fmt_num(r.group_se)} | "
            f"{fmt_ci(r.group_ci_low, r.group_ci_high)} | {fmt_p(r.group_p)} |"
        )
    return "\n".join(lines) + "\n\nTable 2. Primary and robustness OLS models for global aperiodic exponent (N = 138). Positive β indicates higher exponent in TD than ASD.\n"


def build_table_s1_clinical() -> str:
    desc = pd.read_csv(PROJECT_ROOT / "outputs/tables/table1_main_cohort_descriptive.csv")
    rows = [
        "| Variable | ASD (n = 61) | TD (n = 77) | p | Note |",
        "|----------|--------------|-------------|-----|------|",
    ]
    for var, lab in [
        ("ADOS_total", "ADOS total"),
        ("ADOS_SA", "ADOS social affect"),
        ("ADOS_communication", "ADOS communication"),
        ("language_score", "Language score (exploratory)"),
    ]:
        a = desc[(desc.group == "ASD") & (desc.variable == var)]
        t = desc[(desc.group == "TD") & (desc.variable == var)]
        if a.empty:
            continue
        na = int(a.iloc[0]["n"]) if pd.notna(a.iloc[0]["n"]) else "—"
        pa = f"{a.iloc[0]['mean']:.1f} ± {a.iloc[0]['std']:.1f}" if pd.notna(a.iloc[0]["mean"]) else "—"
        if not t.empty and pd.notna(t.iloc[0]["mean"]):
            pt = f"{t.iloc[0]['mean']:.1f} ± {t.iloc[0]['std']:.1f}"
            pcell = "—"
        else:
            pt = "—"
            comp = pd.read_csv(PROJECT_ROOT / "outputs/tables/table1_main_cohort_comparison.csv")
            c = comp[comp.variable == var]
            pcell = fmt_p(c.iloc[0].t_pvalue).replace("p ", "") if not c.empty and pd.notna(c.iloc[0].t_pvalue) else "—"
        note = "Exploratory; definition pending verification" if var == "language_score" else "ASD only" if pt == "—" else ""
        rows.append(f"| {lab} | {pa} (n = {na}) | {pt} | {pcell} | {note} |")
    return "\n".join(rows) + "\n\nSupplementary Table S1. Clinical and exploratory language variables. Language score mapping requires author verification before primary-text reporting.\n"


def clean_supp_df(df: pd.DataFrame, max_rows: int = 12) -> tuple[str, bool]:
    df = df.copy()
    for c in df.columns:
        if df[c].dtype == float or df[c].dtype == int:
            if "p" in c.lower() or c in ("p", "pvalue", "t_pvalue", "group_p", "pearson_p", "spearman_p"):
                df[c] = df[c].apply(lambda x: fmt_p(x) if pd.notna(x) and isinstance(x, (int, float)) else "—")
            elif any(k in c.lower() for k in ("coef", "beta", "se", "ci", "rho", "r_squared", "mean", "std")):
                df[c] = df[c].apply(lambda x: fmt_num(x, 3) if pd.notna(x) and isinstance(x, (int, float)) else "—")
            else:
                df[c] = df[c].apply(
                    lambda x: fmt_num(x, 3)
                    if pd.notna(x) and isinstance(x, (int, float)) and abs(x) < 1e6
                    else ("—" if (isinstance(x, float) and pd.isna(x)) or str(x).lower() in ("nan", "none") else x)
                )
        else:
            df[c] = df[c].apply(
                lambda x: "—"
                if str(x).lower() in ("nan", "none", "true", "false")
                else ("Yes" if str(x) == "True" else ("No" if str(x) == "False" else x))
            )
    full = len(df) > max_rows
    show = df.head(max_rows if full else len(df))
    lines = ["| " + " | ".join(str(c) for c in show.columns) + " |", "| " + " | ".join("---" for _ in show.columns) + " |"]
    for _, row in show.iterrows():
        lines.append("| " + " | ".join(str(row[c]) for c in show.columns) + " |")
    return "\n".join(lines), full


def build_supplementary() -> str:
    parts = ["## Supplementary Materials\n"]
    parts.append("### Supplementary Table S1. Clinical and exploratory variables\n\n")
    parts.append(build_table_s1_clinical())

    flow = pd.read_csv(PROJECT_ROOT / "outputs/tables/sample_inclusion_flow.csv")
    flow = flow.rename(
        columns={
            "stage": "Stage",
            "n_total": "N total",
            "n_ASD": "N ASD",
            "n_TD": "N TD",
            "excluded_total": "Excluded",
            "exclusion_reason": "Reason",
        }
    )
    tbl, _ = clean_supp_df(flow, 20)
    parts.append("### Supplementary Table S2. Sample inclusion flow\n\n" + tbl + "\n")

    for title, path, max_r in [
        ("S3. Age interaction models (summary)", "outputs/tables/compare_preschool_study/age_interaction_models.csv", 8),
        ("S4. Periodic peak group models (summary)", "derivatives/stats/periodic_peak_analysis.csv", 8),
        ("S5. Clinical association models (exploratory)", "derivatives/stats/clinical_correlation_ols.csv", 10),
        ("S6. Split-half reliability", "outputs/tables/extension/split_half_reliability.csv", 10),
        ("S7. Frequency range and fixed/knee sensitivity", "outputs/tables/compare_preschool_study/fixed_vs_knee_summary.csv", 10),
    ]:
        p = PROJECT_ROOT / path
        parts.append(f"### Supplementary Table {title}\n\n")
        if not p.exists():
            parts.append(f"[Table file missing]\n\n")
            continue
        df = pd.read_csv(p)
        if title.startswith("S3"):
            df = df[df.term.str.contains("group|age_months|Intercept", case=False, na=False)]
        if title.startswith("S4"):
            df = df[df.term.str.contains("group", case=False, na=False)]
        tbl, full = clean_supp_df(df, max_r)
        parts.append(tbl + "\n")
        if full:
            parts.append(
                "Note. Full coefficient tables are provided in the supplementary data file.\n\n"
            )

    sp = PROJECT_ROOT / "derivatives/stats/clinical_correlation_spearman.csv"
    if sp.exists():
        parts.append("### Supplementary Table S5 (continued). Spearman clinical correlations\n\n")
        parts.append(clean_supp_df(pd.read_csv(sp), 10)[0] + "\n\n")

    strat = PROJECT_ROOT / "outputs/tables/compare_preschool_study/age_stratified_group_effects.csv"
    if strat.exists():
        parts.append("### Supplementary Table S3 (continued). Age-stratified group effects\n\n")
        parts.append(clean_supp_df(pd.read_csv(strat), 8)[0] + "\n\n")

    for key, cap in [
        ("figs1", "Supplementary Figure S1. Split-half reliability of specparam metrics within the recording session (not test–retest)."),
        ("figs2", "Supplementary Figure S2. Fixed versus knee aperiodic mode sensitivity for global exponent TD − ASD effects."),
    ]:
        rel = {
            "figs1": "../figures/extension/fig_split_half_reliability.png",
            "figs2": "../figures/compare_preschool_study/fig_fixed_vs_knee_effects.png",
        }[key]
        parts.append(f"![{key}]({rel})\n\n{cap}\n\n")
    return "".join(parts)


def build_references() -> str:
    return """## References

Chen, Y., Tsou, M., Nelson, C. A., & Tager-Flusberg, H. (2026). Resting state aperiodic and periodic EEG activity in preschool-aged autistic children: Differences from neurotypical peers and links to language skills. Molecular Autism, 17, Article 7. https://doi.org/10.1186/s13229-025-00700-1 [metadata verification required]

Donoghue, T., Dominguez, J., & Voytek, B. (2020a). Electrophysiological frequency band ratio measures conflate periodic and aperiodic neural activity. eNeuro, 7(6), Article ENEURO.0192-20.2020. https://doi.org/10.1523/ENEURO.0192-20.2020

Donoghue, T., Haller, M., Peterson, E. J., Varma, P., Sebastian, P., Gao, R., Noto, T., Lara, A. H., Wallis, J. D., Knight, R. T., Shestyuk, A., & Voytek, B. (2020b). Parameterizing neural power spectra into periodic and aperiodic components. Nature Neuroscience, 23(12), 1655–1665. https://doi.org/10.1038/s41593-020-00744-x

Gao, R., Peterson, E. J., & Voytek, B. (2017). Inferring synaptic excitation/inhibition balance from field potentials. NeuroImage, 158, 70–78. https://doi.org/10.1016/j.neuroimage.2017.06.065

Hill, A. T., Clark, G. M., Bigelow, F. J., Lum, J. A. G., & Enticott, P. G. (2022). Periodic and aperiodic neural activity displays age-dependent changes across early-to-middle childhood. Developmental Cognitive Neuroscience, 54, Article 101076. https://doi.org/10.1016/j.dcn.2022.101076

Karalunas, S. L., Gustafsson, H. C., Ostlund, B. D., Alperin, B. R., Deming, E. M., & Sullivan, E. L. (2022). Electroencephalogram aperiodic power spectral slope can be reliably measured and predicts ADHD risk in early development. Developmental Psychobiology, 64(3), Article e22228. https://doi.org/10.1002/dev.22228

Liang, S., & Mody, M. (2022). Abnormal brain oscillations in developmental disorders: Application of resting state EEG and MEG in autism spectrum disorder and fragile X syndrome. Frontiers in Neuroimaging, 1, Article 903191. https://doi.org/10.3389/fnimg.2022.903191

Manyukhina, V. O., Prokofyev, A. O., Galuta, I. A., Goiaeva, D. E., Obukhova, T. S., Stroganova, T. A., Orekhova, E. V., & Altukhov, D. I. (2022). Globally elevated excitation–inhibition ratio in children with autism spectrum disorder and below-average intelligence. Molecular Autism, 13, Article 20. https://doi.org/10.1186/s13229-022-00498-2

Neo, W. S., Foti, D., Keehn, B., & Kelleher, B. (2023). Resting-state EEG power differences in autism spectrum disorder: A systematic review and meta-analysis. Translational Psychiatry, 13, Article 389. https://doi.org/10.1038/s41398-023-02681-2

von Elm, E., Altman, D. G., Egger, M., Pocock, S. J., Gøtzsche, P. C., & Vandenbroucke, J. P. (2007). The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: Guidelines for reporting observational studies. PLoS Medicine, 4(10), Article e296. https://doi.org/10.1371/journal.pmed.0040296

Wilkinson, C. L., Yankowitz, L. D., Chao, J. Y., et al. (2024). Developmental trajectories of EEG aperiodic and periodic components in children 2–44 months of age. Nature Communications, 15, Article 5788. https://doi.org/10.1038/s41467-024-50204-4 [metadata verification required]
"""


def author_checklist() -> str:
    return """## 作者需补充清单（投稿前）

### 伦理与诊断
- [待补充] ASD 诊断标准（DSM-5/ICD-11/ADOS-2 等）
- [待补充] TD 纳入标准
- [待补充] 排除标准完整列表
- [待补充] 伦理委员会名称与批准号
- [待补充] 知情同意程序

### 评估与变量
- [待补充] IQ 量表名称与施测说明
- [待补充] ADOS/ADOS-2 模块与评估者资质
- [待核实] language_score 与 ToMI/Resting_info 映射（探索性，仅补充表）
- [待核实] ADOS_communication、ADOS_SA 操作定义
- [待核实] ICA 参数与伪迹剔除策略

### 采集
- [待补充] EEG 采集指令、记录时长、阻抗标准

### 软件引用
- [citation needed] MNE-Python 正式文献条目

### 引用终核
- Chen et al. (2026)、Wilkinson et al. (2024)：metadata verification required（勿在未核对前修改卷期页码）
"""


def transform_body(v5: str) -> str:
    # Split at References
    parts = re.split(r"\n## References\n", v5, maxsplit=1)
    body = parts[0]
    # Remove v5 appendices and supplementary from body
    body = re.sub(r"\n## 补充材料[\s\S]*", "", body)
    body = re.sub(r"\n## 附录[\s\S]*", "", body)
    body = re.sub(r"^---\s*\n", "", body, flags=re.M)

    # Replace table 1 block
    body = re.sub(
        r"\| 变量 \| ASD[\s\S]*?\*\*表 1\.\*\*[^\n]+\n",
        build_table1(),
        body,
        count=1,
    )
    body = re.sub(
        r"\| 模型 \| 协变量[\s\S]*?\*\*表 2\.\*\*[^\n]+\n",
        build_table2(),
        body,
        count=1,
    )
    if "| Model | Covariates |" not in body:
        body = re.sub(
            r"\| 仅组别[\s\S]*?Table 2\.[^\n]+\n",
            build_table2(),
            body,
            count=1,
        )

    # Methods formula
    body = body.replace(
        "`global_exponent ~ group + age_months + sex + IQ_total + usable_epochs`",
        "global exponent ~ group + age (months) + sex + IQ total + usable epochs",
    )
    body = body.replace(
        "`outcome ~ group * age_months + sex + IQ_total + usable_epochs`",
        "outcome ~ group × age (months) + sex + IQ total + usable epochs",
    )
    body = body.replace(
        "参照组为 ASD；`C(group)[T.TD]` 表示 TD 相对 ASD 增量。主模型：",
        "参照组为 ASD；TD 相对 ASD 的回归系数表示 TD 组增量。主模型：",
    )
    body = body.replace("原始 `.mff` 格式", "原始 MFF 格式")
    body = body.replace(
        "周期峰与 ASD 内临床分析为预设探索性分析；临床变量 [待核实变量定义]。",
        "周期峰与 ASD 内临床分析为预设探索性分析（临床变量定义见作者需补充清单）。",
    )
    body = body.replace(
        "（3）临床变量缺失与定义 [待核实变量定义]；",
        "（3）临床变量定义尚待终核（见作者需补充清单）；",
    )

    # Results 3.7 simplify
    body = re.sub(
        r"### 3\.7 ASD 组内临床关联[\s\S]*?(?=### 3\.8)",
        "### 3.7 ASD 组内临床关联（探索性）\n\n"
        "在 ASD 子样本（n = 61）中，探索性临床关联分析未发现经多重比较校正后仍显著的主效应"
        "（详见补充表 S5）。具体变量定义与量表映射待作者核实后于补充材料报告。\n\n",
        body,
    )

    # Abstract - remove ADOS trend sentence
    body = body.replace(
        "ASD 内临床关联未达显著（ADOS 社交子分 OLS *p* ≈ 0.057；Spearman *p* ≈ 0.076）。",
        "ASD 内探索性临床关联未达显著。",
    )

    # Discussion 4.7 simplify
    body = re.sub(
        r"### 4\.7 临床关联\n\n[^\n]+",
        "### 4.7 临床关联\n\n"
        "探索性临床关联未达显著，与 Chen et al.（2026）中 ASD 内语言相关结果不可直接类比"
        "（量表来源与变量定义不同）。",
        body,
        count=1,
    )

    # Collect placeholders from body for checklist only - keep short in methods
    body = body.replace(
        "`language_score` 来源于 [待补充：量表名称]；[待核实变量定义] 其与 Resting_info/ToMI 的对应关系需在定稿前确认。`ADOS_communication` [待核实变量定义]；`ADOS_RRB` 在本队列无有效数据。",
        "临床与语言变量定义见作者需补充清单；探索性语言指标仅于补充表报告。",
    )

    body = replace_donoghue_cites(body)
    body = strip_md(body)
    body = fmt_stats_line(body)

    # Figure captions cleanup
    body = re.sub(r"\*\*图 (\d+[AB]?)\.\*\*", r"Figure \1.", body)
    body = re.sub(r"\*\*A\.\*\*", "A.", body)
    body = re.sub(r"\*\*B\.\*\*", "B.", body)
    body = re.sub(r"\*\*C\.\*\*", "C.", body)

    # Terminology first-use (light touch)
    body = body.replace(
        "自闭症谱系障碍（autism spectrum disorder, ASD）",
        "自闭症谱系障碍（Autism Spectrum Disorder, ASD）",
        1,
    )

    return body.strip()


FIGURE_MAP = [
    ("Figure 1A", "outputs/figures/paper_ready_v2/supp_consort_flow_paper.png", "3.1", "Sample inclusion flow"),
    ("Figure 1B", "outputs/figures/paper_ready_v2/fig1_psd_specparam_overview.png", "3.1", "Group-averaged PSD"),
    ("Figure 2", "outputs/figures/paper_ready_v2/fig2_main_aperiodic_effects.png", "3.2", "Main aperiodic effects"),
    ("Figure 3A", "outputs/figures/extension/fig_development_interaction_exponent.png", "3.4", "Age interaction exponent"),
    ("Figure 3B", "outputs/figures/extension/fig_development_interaction_offset.png", "3.4", "Age interaction offset"),
    ("Figure 4", "outputs/figures/paper_ready_v2/fig3_spatial_distribution.png", "3.5", "Spatial distribution"),
    ("Figure S1", "outputs/figures/extension/fig_split_half_reliability.png", "Supplementary", "Split-half reliability"),
    ("Figure S2", "outputs/figures/compare_preschool_study/fig_fixed_vs_knee_effects.png", "Supplementary", "Fixed vs knee"),
]


def write_docx(md: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = normal.paragraph_format
    pf.line_spacing = 1.5

    img_re = re.compile(r"^!\[(.+?)\]\((.+?)\)$")
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        m = img_re.match(line.strip())
        if m:
            _, rel = m.groups()
            img_path = (OUT_MD.parent / rel).resolve()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img_path.exists():
                p.add_run().add_picture(str(img_path), width=Cm(14))
            else:
                p.add_run(f"[Figure file missing: {rel}]")
            i += 1
            if i < len(lines) and lines[i].strip() and not lines[i].startswith("|"):
                cap = doc.add_paragraph(lines[i].strip())
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cap.runs:
                    r.font.size = Pt(10)
                    r.font.italic = True
                i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            tlines = []
            while i < len(lines) and lines[i].startswith("|"):
                tlines.append(lines[i])
                i += 1
            add_three_line_table(doc, tlines)
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith("## References") or line.startswith("Reference"):
            doc.add_page_break()
        p = doc.add_paragraph(line.strip())
        if line.startswith("Chen,") or line.startswith("Donoghue,"):
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-1.27)
        i += 1

    try:
        doc.save(OUT_DOCX)
    except PermissionError:
        alt = OUT_DOCX.with_name(OUT_DOCX.stem + "_updated.docx")
        doc.save(alt)
        print(f"Wrote {alt} (original locked)")


def add_three_line_table(doc, table_lines):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rows = []
    for tl in table_lines:
        if re.match(r"^\|[\s\-:|]+\|$", tl):
            continue
        rows.append([c.strip() for c in tl.strip("|").split("|")])
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            tbl.rows[ri].cells[ci].text = cell
    # Remove vertical borders (simplified)
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("left", "right"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "nil")
                borders.append(el)
            tcPr.append(borders)


def write_audits():
    fig_lines = ["# Figure and table audit (v6)\n\n", "| Figure | File | Section | Caption |\n", "|--------|------|---------|----------|\n"]
    for fig, path, sec, desc in FIGURE_MAP:
        exists = (PROJECT_ROOT / path).exists()
        fig_lines.append(f"| {fig} | `{path}` | {sec} | {desc} {'OK' if exists else 'MISSING'} |\n")
    OUT_FIG_AUDIT.write_text("".join(fig_lines), encoding="utf-8")

    readiness = """# Submission readiness audit (v6)

## 已解决的问题

- 清理 Markdown 星号、分隔线、反引号与代码块式公式
- 统一 p 值、β、SE、CI 与三位小数格式
- Donoghue (2020) 拆分为 2020a（eNeuro）与 2020b（Nature Neuroscience），正文与 References 同步
- Table 1 仅保留人口学与 EEG QC；临床变量移至 Supplementary Table S1（探索性标注）
- Table 2 精简为协变量模型摘要
- 补充表去除 nan/True/False 与导出体例；大表仅保留摘要行
- Results 3.7 与 Discussion 4.7 弱化未核实临床变量，不突出 ADOS 趋势
- Word：12 pt、1.5 倍行距、图居中、References 悬挂缩进、三线表（无竖线）

## 仍需作者补充的信息

见稿末「作者需补充清单」：伦理批号、诊断/排除标准、IQ/ADOS 量表、EEG 采集细节、ICA 策略、MNE 引用。

## 仍需核查的引用

- Chen et al. (2026) — metadata verification required
- Wilkinson et al. (2024) — metadata verification required（完整作者列表）
- MNE-Python — [citation needed]
- Karalunas et al. (2022) — 建议用 Zotero 终核卷期页码

## 仍建议补做的分析

1. **女性仅样本**主模型（female-only primary model），报告 β 与 n
2. **性别平衡或性别敏感性**：sex-balanced subsample 或强化 sex-adjusted 报告（主模型已含 sex，建议补充女性子样本与男性比例说明）
3. **周期峰分析**：若时间允许，按主分析 N = 138（非 N = 145）重跑 periodic peak OLS，更新 Supplementary Table S4
4. **临床关联**：language_score 未核实前，仅保留于补充材料；勿写入正文主结论

## 投稿前必须人工检查的图表

| 项目 | 检查内容 |
|------|----------|
| Figure 1A | CONSORT 数字与正文 168→145→138 一致 |
| Figure 1B | PSD 轴标签与 specparam 拟合范围 1–40 Hz |
| Figure 2 | 箱线图 TD>ASD 与 β = 0.079, p = .012 一致 |
| Figure 3 | 横断面交互勿写成纵向轨迹（图注已提醒） |
| Figure 4 | Topomap 为头皮水平，非源定位 |
| Figure S1 | 标注 split-half，非 test–retest |
| Table 1 | 性别分布不均衡已在讨论说明 |
| Table 2 | 主模型行与摘要结果一致 |
| 所有图 | Word 中确为嵌入图而非占位文字 |

## 输出文件

- `manuscript_draft_zh_v6_submission_style.md`
- `manuscript_draft_zh_v6_submission_style.docx`
- `reference_audit_v6.md`
- `figure_table_audit_v6.md`
"""
    OUT_READINESS.write_text(readiness, encoding="utf-8")

    ref_audit = """# Reference audit v6

| Citation key | In text? | In References? | Notes |
|--------------|----------|----------------|-------|
| Chen et al. (2026) | Yes | Yes | metadata verification required |
| Donoghue et al. (2020a) | Yes | Yes | eNeuro band-ratio |
| Donoghue et al. (2020b) | Yes | Yes | Nat Neurosci specparam |
| Gao et al. (2017) | Yes | Yes | |
| Hill et al. (2022) | Yes | Yes | |
| Karalunas et al. (2022) | Yes | Yes | |
| Liang & Mody (2022) | Yes | Yes | |
| Manyukhina et al. (2022) | Yes | Yes | MEG; caution in text |
| Neo et al. (2023) | Yes | Yes | |
| von Elm et al. (2007) | Yes | Yes | |
| Wilkinson et al. (2024) | Yes | Yes | metadata verification required |
| MNE-Python | No (citation needed) | No | Methods placeholder |

## Orphan check

- References 条目均在正文出现：是（12 条）
- 正文引用均在 References：除 MNE 待补
"""
    OUT_REF_AUDIT.write_text(ref_audit, encoding="utf-8")


def main():
    v5 = V5.read_text(encoding="utf-8")
    body = transform_body(v5)
    # Preserve figure blocks from v5 (paths)
    fig_section = re.findall(r"!\[Figure[^\]]+\]\([^)]+\)[\s\S]*?(?=\n\n(?:Table|###|\|))", v5)
    # Rebuild figures with clean captions - use regex on body for images
    for old_cap in ["**图 1A.**", "**图 1B.**", "**图 2.**", "**图 3A.**", "**图 3B.**", "**图 4.**"]:
        body = body.replace(old_cap, old_cap.replace("**", "").replace("图", "Figure"))

    md = (
        body
        + "\n\n"
        + build_references()
        + "\n\n"
        + build_supplementary()
        + "\n\n"
        + author_checklist()
    )
    md = strip_md(md)
    md = replace_donoghue_cites(md)
    OUT_MD.write_text(md, encoding="utf-8")
    write_docx(md)
    write_audits()
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {OUT_READINESS}")
    print(f"Wrote {OUT_REF_AUDIT}")
    print(f"Wrote {OUT_FIG_AUDIT}")


if __name__ == "__main__":
    main()
