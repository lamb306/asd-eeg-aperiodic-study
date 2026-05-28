#!/usr/bin/env python
"""Build manuscript v5 with embedded figures/tables (Markdown + DOCX)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

import pandas as pd

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

V4 = PROJECT_ROOT / "outputs/reports/manuscript_draft_zh_v4_APA.md"
OUT_MD = PROJECT_ROOT / "outputs/reports/manuscript_draft_zh_v5_APA_figures.md"
OUT_DOCX = PROJECT_ROOT / "outputs/reports/manuscript_draft_zh_v5_APA_figures.docx"
OUT_SUPP_DOCX = PROJECT_ROOT / "outputs/reports/supplementary_materials_zh.docx"
OUT_AUDIT = PROJECT_ROOT / "outputs/reports/figure_table_insertion_audit.md"
PNG_CACHE = PROJECT_ROOT / "outputs/figures/_png_cache_v5"

# figure key -> (relative path from project root, caption block)
FIGURES = {
    "fig1a": (
        "outputs/figures/paper_ready_v2/supp_consort_flow_paper.svg",
        None,
    ),
    "fig1b": (
        "outputs/figures/paper_ready_v2/fig1_psd_specparam_overview.svg",
        None,
    ),
    "fig2": (
        "outputs/figures/paper_ready_v2/fig2_main_aperiodic_effects.svg",
        None,
    ),
    "fig3a": (
        "outputs/figures/extension/fig_development_interaction_exponent.svg",
        None,
    ),
    "fig3b": (
        "outputs/figures/extension/fig_development_interaction_offset.svg",
        None,
    ),
    "fig4": (
        "outputs/figures/paper_ready_v2/fig3_spatial_distribution.svg",
        None,
    ),
    "figs1": (
        "outputs/figures/extension/fig_split_half_reliability.svg",
        None,
    ),
    "figs2": (
        "outputs/figures/compare_preschool_study/fig_fixed_vs_knee_effects.svg",
        None,
    ),
}

CAPTION_FIG1 = """**图 1.** 样本纳入流程与组平均功率谱密度（PSD）。

**A.** CONSORT 式流程图，展示静息态 EEG 非周期分析的被试纳入过程。168 名被试完成预处理；145 名达到至少 60 个 2 s 可用 epoch；另有 7 名因超过 20% 通道未通过 specparam 质量控制而排除，主分析队列 138 名（ASD 61，TD 77）。

**B.** ASD 与 TD 组平均 PSD（specparam 拟合频段 1–40 Hz）。阴影为跨被试标准误。"""

CAPTION_FIG2 = """**图 2.** 全局非周期参数的组间差异与稳健性分析。

**A.** 按组显示的 global aperiodic exponent（箱线图叠加个体散点）。标注为校正年龄、性别、IQ_total 与可用 epoch 数后的组效应估计。

**B.** 按组显示的 global aperiodic offset。

**C.** global exponent 组效应（TD − ASD）在嵌套 OLS 模型中的稳健性；点为 β，横线为 95% CI。"""

CAPTION_FIG3 = """**图 3.** 非周期 EEG 参数的年龄依赖性组效应。

**A.** global aperiodic exponent 随年龄的模型预测关联（按组）。

**B.** global aperiodic offset 随年龄的模型预测关联（按组）。

group × age 交互对 exponent 与 offset 均显著。本研究为横断面设计，**不应**将上述模式解释为个体纵向发育轨迹。"""

CAPTION_FIG4 = """**图 4.** aperiodic exponent 组效应的空间分布（探索性）。

**A.** 各 ROI 上 TD − ASD 对 exponent 的边际效应（混合模型）。

**B.** 64 通道 OLS 回归系数头皮拓扑图（HydroCel-64 布局）；FDR *q* < .05 的通道以标记显示（E33、E36、E37、E38）。**仅反映头皮电极水平，不作皮层源定位推断。**"""

CAPTION_FIGS1 = """**图 S1.** specparam 衍生指标的奇偶 epoch 分半信度。

global exponent、global offset 与 alpha 峰功率的组内分半 Spearman 相关及 Spearman–Brown 校正估计均 > 0.97。**为同次记录内的 within-session 一致性，不是跨日 test–retest 重测信度。**"""

CAPTION_FIGS2 = """**图 S2.** fixed 与 knee 模式的敏感性分析。

fixed 模式各频段范围的 TD − ASD exponent 效应方向一致且多数达到 *p* < .05；knee 模式方向相似但统计证据较弱。"""


def svg_to_png(svg_path: Path, png_path: Path, scale: float = 2.0) -> bool:
    if png_path.exists() and png_path.stat().st_mtime >= svg_path.stat().st_mtime:
        return True
    try:
        import cairosvg

        png_path.parent.mkdir(parents=True, exist_ok=True)
        cairosvg.svg2png(
            url=str(svg_path.resolve()),
            write_to=str(png_path),
            scale=scale,
        )
        return True
    except Exception as e:
        print(f"PNG convert failed {svg_path}: {e}")
        return False


def resolve_figure(rel: str) -> tuple[Path | None, list[str]]:
    """Return best path (png preferred) and checked candidates."""
    root = PROJECT_ROOT
    base = Path(rel)
    stem = base.stem
    parent = base.parent
    checked = []
    for ext in (".png", ".svg", ".pdf"):
        p = root / parent / f"{stem}{ext}"
        checked.append(str(p))
        if p.exists() and ext == ".png":
            return p, checked
    svg = root / base
    if svg.suffix == ".svg" and svg.exists():
        png_side = root / parent / f"{stem}.png"
        png_cache = PNG_CACHE / parent.name / f"{stem}.png"
        for png_out in (png_side, png_cache):
            if svg_to_png(svg, png_out):
                return png_out, checked
    return None, checked


def md_img(rel_from_reports: str, alt: str) -> str:
    return f"![{alt}]({rel_from_reports})"


def build_table1_md() -> str:
    desc = pd.read_csv(PROJECT_ROOT / "outputs/tables/table1_main_cohort_descriptive.csv")
    comp = pd.read_csv(PROJECT_ROOT / "outputs/tables/table1_main_cohort_comparison.csv")
    qc = pd.read_csv(PROJECT_ROOT / "outputs/tables/table2_main_cohort_eeg_qc.csv")
    cohort = pd.read_csv(PROJECT_ROOT / "outputs/tables/main_cohort_subject_list.csv")
    deriv = PROJECT_ROOT / "derivatives"
    preproc = deriv / "qc" / "preproc_summary.csv"
    sp = deriv / "specparam" / "specparam_qc_summary_subject.csv"
    if preproc.exists():
        cohort = cohort.merge(
            pd.read_csv(preproc)[["subject_id", "bad_channel_count"]],
            on="subject_id",
            how="left",
            suffixes=("", "_pre"),
        )
    if sp.exists():
        spdf = pd.read_csv(sp)[["subject_id", "mean_r_squared"]]
        if "mean_r_squared" in cohort.columns:
            cohort = cohort.drop(columns=["mean_r_squared"])
        cohort = cohort.merge(spdf, on="subject_id", how="left")

    def row_desc(var, label, decimals=1):
        a = desc[(desc.group == "ASD") & (desc.variable == var)]
        t = desc[(desc.group == "TD") & (desc.variable == var)]
        c = comp[comp.variable == var]
        if a.empty or t.empty:
            return None
        d = decimals
        ma, sa = a.iloc[0]["mean"], a.iloc[0]["std"]
        mt, st = t.iloc[0]["mean"], t.iloc[0]["std"]
        pa = f"{ma:.{d}f} ± {sa:.{d}f}" if pd.notna(ma) else "—"
        pt = f"{mt:.{d}f} ± {st:.{d}f}" if pd.notna(mt) else "—"
        p = ""
        if not c.empty and pd.notna(c.iloc[0].t_pvalue):
            pv = c.iloc[0].t_pvalue
            p = "< 0.001" if pv < 0.001 else f"{pv:.3f}"
        return f"| {label} | {pa} | {pt} | {p} |"

    rows = [
        "| 变量 | ASD (*n* = 61) | TD (*n* = 77) | *p* |",
        "|------|----------------|---------------|-----|",
        "| *n* | 61 | 77 | — |",
    ]
    for var, label in [
        ("age_months", "年龄（月）"),
        ("IQ_total", "IQ 总分"),
    ]:
        r = row_desc(var, label)
        if r:
            rows.append(r)

    if "sex" in cohort.columns:
        ct = pd.crosstab(cohort["group"], cohort["sex"])
        from scipy.stats import chi2_contingency

        _, p_sex, _, _ = chi2_contingency(ct.values)
        def sex_str(g):
            sub = cohort[cohort.group == g]
            f = int((sub.sex == "F").sum())
            m = int((sub.sex == "M").sum())
            return f"{int(f)} 女 / {int(m)} 男"
        pstr = "< 0.001" if p_sex < 0.001 else f"{p_sex:.3f}"
        rows.append(f"| 性别 | {sex_str('ASD')} | {sex_str('TD')} | {pstr}† |")

    rows.append("| ADOS 总分‡ | 14.1 ± 3.1 | — | — |")
    rows.append("| ADOS 社交（`ADOS_SA`）‡ | 9.3 ± 2.0 | — | — |")
    rows.append("| ADOS 沟通‡ | 4.9 ± 1.3 | — | — |")
    r = row_desc("language_score", "语言分数（ToMI）§")
    if r:
        rows.append(r.replace("language_score", "语言分数"))

    qc_map = {
        "usable_epochs": ("可用 epoch 数", 1),
        "usable_seconds": ("可用记录时长（s）", 1),
        "bad_channel_count": ("坏导数量", 1),
        "mean_r_squared": ("平均拟合 *R*²", 3),
    }
    if "usable_epochs" in cohort.columns and "usable_seconds" not in cohort.columns:
        cohort = cohort.copy()
        cohort["usable_seconds"] = cohort["usable_epochs"] * 2.0
    for var, spec in qc_map.items():
        label, dec = spec if isinstance(spec, tuple) else (spec, 1)
        c = qc[qc.variable == var]
        if c.empty:
            continue
        row = c.iloc[0]
        if var in cohort.columns:
            sa = cohort.loc[cohort.group == "ASD", var]
            st = cohort.loc[cohort.group == "TD", var]
            pa = f"{sa.mean():.{dec}f} ± {sa.std():.{dec}f}"
            pt = f"{st.mean():.{dec}f} ± {st.std():.{dec}f}"
        else:
            pa = f"{row.mean_a:.{dec}f}"
            pt = f"{row.mean_b:.{dec}f}"
        pv = row.t_pvalue
        p = "< 0.001" if pd.notna(pv) and pv < 0.001 else (f"{pv:.3f}" if pd.notna(pv) else "")
        rows.append(f"| {label} | {pa} | {pt} | {p} |")

    note = """
† 卡方检验。‡ 仅 ASD。§ `language_score` 映射 ToMI；ASD 缺失 6 例。

**表 1.** 主分析队列（*N* = 138）人口学、临床特征与 EEG/specparam 质量控制摘要。连续变量为均值 ± 标准差；组间检验为独立样本 *t* 检验（双侧），除非注明。
"""
    return "\n".join(rows) + "\n" + note


def build_table2_md() -> str:
    df = pd.read_csv(PROJECT_ROOT / "outputs/tables/global_exponent_robustness_models.csv")
    labels = {
        "model_1": "仅组别",
        "model_2": "+ 年龄、性别",
        "model_3": "+ IQ_total",
        "model_4": "主模型（+ 可用 epoch）",
        "model_5": "主模型 + 平均 *R*²",
        "model_6": "主模型 + 坏导数",
    }
    lines = [
        "| 模型 | 协变量 | *N* | β (TD − ASD) | SE | 95% CI | *p* |",
        "|------|--------|-----|--------------|-----|--------|-----|",
    ]
    for _, r in df.iterrows():
        name = r["model_name"]
        if name == "model_7_exclude_IQ_below_70":
            continue
        lab = labels.get(name, name)
        ci = f"[{r['group_ci_low']:.3f}, {r['group_ci_high']:.3f}]"
        p = r["group_p"]
        ps = "< 0.001" if p < 0.001 else f"{p:.3f}"
        lines.append(
            f"| {lab} | {r['covariates']} | {int(r['n'])} | "
            f"{r['group_coef_TD_vs_ASD']:.3f} | {r['group_se']:.3f} | {ci} | {ps} |"
        )
    note = "\n\n**表 2.** global aperiodic exponent 的主模型与稳健性 OLS 模型。β > 0 表示 TD > ASD。\n"
    return "\n".join(lines) + note


def df_to_md(df: pd.DataFrame, max_rows: int | None = None) -> str:
    if max_rows:
        df = df.head(max_rows)
    cols = list(df.columns)
    lines = ["| " + " | ".join(str(c) for c in cols) + " |", "| " + " | ".join("---" for _ in cols) + " |"]
    for _, row in df.iterrows():
        lines.append("| " + " | ".join(str(row[c]) for c in cols) + " |")
    return "\n".join(lines)


def rel_md(path: Path) -> str:
    """Path relative to outputs/reports for markdown."""
    return "../" + path.relative_to(PROJECT_ROOT / "outputs").as_posix()


def figure_block(key: str, label: str, caption: str, audit: dict) -> str:
    rel, _ = FIGURES[key]
    resolved, checked = resolve_figure(rel)
    audit["figures"].append(
        {
            "key": key,
            "label": label,
            "resolved": str(resolved) if resolved else None,
            "checked": checked,
        }
    )
    if resolved is None:
        audit["missing_figures"].append((label, rel, checked))
        return f"\n[Figure file missing: {Path(rel).name}]\n\n{caption}\n"
    rmd = rel_md(resolved)
    return f"\n{md_img(rmd, label)}\n\n{caption}\n"


def insert_after_section(text: str, section_header: str, block: str) -> str:
    """Insert block after the paragraph following section_header (first blank line after section)."""
    idx = text.find(section_header)
    if idx < 0:
        return text
    start = idx + len(section_header)
    # find next ### or --- after this section's content until next subsection
    next_sec = re.search(r"\n### 3\.\d", text[start:])
    if not next_sec:
        next_sec = re.search(r"\n---\n", text[start:])
    if next_sec:
        end = start + next_sec.start()
    else:
        end = len(text)
    return text[:end] + block + text[end:]


def build_supplementary_md() -> str:
    parts = ["\n---\n\n## 补充材料\n"]
    supp_tables = [
        (
            "表 S1",
            "样本纳入流程",
            "outputs/tables/sample_inclusion_flow.csv",
            True,
        ),
        (
            "表 S2",
            "拟合频段与 fixed/knee 敏感性分析",
            "outputs/tables/compare_preschool_study/fixed_vs_knee_summary.csv",
            True,
        ),
        (
            "表 S3",
            "年龄交互与分层组效应",
            "outputs/tables/compare_preschool_study/age_interaction_models.csv",
            False,
        ),
        (
            "表 S4",
            "周期峰参数分析",
            "derivatives/stats/periodic_peak_analysis.csv",
            False,
        ),
        (
            "表 S5",
            "临床关联分析",
            "derivatives/stats/clinical_correlation_ols.csv",
            False,
        ),
        (
            "表 S6",
            "奇偶分半信度",
            "outputs/tables/extension/split_half_reliability.csv",
            True,
        ),
    ]
    for title, desc, rel, embed_small in supp_tables:
        p = PROJECT_ROOT / rel
        parts.append(f"\n### {title}. {desc}\n\n")
        if not p.exists():
            parts.append(f"[Table file missing: {rel}]\n")
            continue
        df = pd.read_csv(p)
        if embed_small and len(df) <= 12:
            parts.append(df.to_markdown(index=False))
            parts.append("\n")
        else:
            parts.append(
                f"完整数据见项目补充表（{len(df)} 行）。以下为摘要：\n\n"
            )
            parts.append(df.head(8).to_markdown(index=False))
            parts.append("\n\n*（完整表见补充数据包。）*\n")
    parts.append(
        figure_block("figs1", "Figure S1", CAPTION_FIGS1, {"figures": [], "missing_figures": []})
    )
    # fix figure_block side effect - will handle in main
    return "".join(parts)


def main() -> None:
    audit = {
        "figures": [],
        "missing_figures": [],
        "tables": [],
        "missing_tables": [],
    }
    text = V4.read_text(encoding="utf-8")
    # strip old appendix figure list - keep 待补充 and 引用待核查, update 图表清单 later
    split_refs = text.split("## References")
    body = split_refs[0]
    refs_tail = "## References" + split_refs[1] if len(split_refs) > 1 else ""

    # Remove v4 附录：图表清单 through 生成统计
    body = re.sub(
        r"\n---\n\n## 附录：图表清单[\s\S]*?(?=\n---\n\n## 附录：待补充|\Z)",
        "\n",
        body,
        count=1,
    )
    body = re.sub(r"\n---\n\n## 生成统计[\s\S]*$", "", body)

    # §3.1 + figures + table1
    b1 = (
        figure_block("fig1a", "Figure 1A", "**图 1A.** 样本纳入流程。", audit)
        + figure_block("fig1b", "Figure 1B", "**图 1B.** 组平均 PSD 与 specparam 拟合示意。", audit)
        + "\n" + CAPTION_FIG1 + "\n"
        + "\n" + build_table1_md()
    )
    body = insert_after_section(body, "### 3.1 样本纳入与 EEG 质量控制\n", b1)

    # Update 3.1 text with stats from table1
    body = body.replace(
        "两组年龄差异不显著 [待补充具体统计量]。TD 组 IQ_total 高于 ASD [待补充具体统计量]",
        "两组年龄差异不显著（*p* = 0.319）。TD 组 IQ_total 高于 ASD（*p* < 0.001）",
        1,
    )

    b2 = figure_block("fig2", "Figure 2", CAPTION_FIG2, audit)
    body = insert_after_section(body, "### 3.2 Global aperiodic exponent 与 offset\n", b2)

    b3 = "\n" + build_table2_md()
    body = insert_after_section(body, "### 3.3 稳健性与敏感性\n", b3)

    b4 = (
        figure_block("fig3a", "Figure 3A", "**图 3A.** exponent。", audit)
        + figure_block("fig3b", "Figure 3B", "**图 3B.** offset。", audit)
        + "\n" + CAPTION_FIG3 + "\n"
    )
    body = insert_after_section(body, "### 3.4 年龄依赖性\n", b4)

    b5 = figure_block("fig4", "Figure 4", CAPTION_FIG4, audit)
    body = insert_after_section(body, "### 3.5 空间分布（探索性）\n", b5)

    body = body.replace(
        "### 3.8 Split-half 信度\n\n*n* = 138。",
        "### 3.8 Split-half 信度\n\n*n* = 138（见图 S1）。",
        1,
    )

    # Supplementary after references
    supp_parts = ["\n---\n\n## 补充材料\n\n"]
    supp_audit_tables = []
    for title, desc, rel, embed in [
        ("S1", "样本纳入流程", "outputs/tables/sample_inclusion_flow.csv", True),
        ("S2", "拟合频段与 fixed/knee 敏感性", "outputs/tables/compare_preschool_study/fixed_vs_knee_summary.csv", True),
        ("S3", "年龄交互与分层组效应", "outputs/tables/compare_preschool_study/age_interaction_models.csv", False),
        ("S4", "周期峰参数分析", "derivatives/stats/periodic_peak_analysis.csv", False),
        ("S5", "临床关联（OLS 与 Spearman）", "derivatives/stats/clinical_correlation_ols.csv", False),
        ("S6", "奇偶分半信度", "outputs/tables/extension/split_half_reliability.csv", True),
    ]:
        p = PROJECT_ROOT / rel
        supp_parts.append(f"### 表 {title}. {desc}\n\n")
        if not p.exists():
            supp_parts.append(f"[Table file missing: {rel}]\n\n")
            audit["missing_tables"].append((title, rel))
            continue
        df = pd.read_csv(p)
        audit["tables"].append((title, rel, len(df), "embedded" if embed and len(df) <= 15 else "summary"))
        if embed and len(df) <= 15:
            supp_parts.append(df_to_md(df) + "\n\n")
        else:
            supp_parts.append(f"*完整表共 {len(df)} 行；定稿时可从补充数据包导出。*\n\n")
            supp_parts.append(df_to_md(df, max_rows=6) + "\n\n")

    supp_parts.append(figure_block("figs1", "Figure S1", CAPTION_FIGS1, audit))
    supp_parts.append(figure_block("figs2", "Figure S2", CAPTION_FIGS2, audit))

    # S3 second file note
    p3b = PROJECT_ROOT / "outputs/tables/compare_preschool_study/age_stratified_group_effects.csv"
    if p3b.exists():
        supp_parts.append("### 表 S3（续）. 年龄分层组效应\n\n")
        supp_parts.append(df_to_md(pd.read_csv(p3b), max_rows=8) + "\n\n")
        audit["tables"].append(("S3b", str(p3b.relative_to(PROJECT_ROOT)), len(pd.read_csv(p3b)), "summary"))
    sp5 = PROJECT_ROOT / "derivatives/stats/clinical_correlation_spearman.csv"
    if sp5.exists():
        supp_parts.append("### 表 S5（续）. 临床关联 Spearman\n\n")
        supp_parts.append(df_to_md(pd.read_csv(sp5)) + "\n\n")
        audit["tables"].append(("S5b", str(sp5.relative_to(PROJECT_ROOT)), 4, "embedded"))
    s2b = PROJECT_ROOT / "outputs/tables/sensitivity_analysis_final.csv"
    if s2b.exists():
        supp_parts.append("### 表 S2（续）. 敏感性分析（完整）\n\n")
        supp_parts.append(
            "*完整表共 "
            + str(len(pd.read_csv(s2b)))
            + " 行；摘要见表 S2 主表。*\n\n"
        )

    appendix = """
---

## 附录：图表清单

| 编号 | 内容 | 位置 |
|------|------|------|
| 图 1 | 样本纳入 + 组平均 PSD | 结果 §3.1 |
| 表 1 | 人口学与 EEG QC | 结果 §3.1 |
| 图 2 | 主效应与稳健性 | 结果 §3.2–3.3 |
| 表 2 | 主模型与稳健性模型 | 结果 §3.3 |
| 图 3 | 年龄交互 | 结果 §3.4 |
| 图 4 | 空间分布 | 结果 §3.5 |
| 图 S1 | Split-half 信度 | 补充材料 |
| 图 S2 | fixed vs knee | 补充材料 |
| 表 S1–S6 | 见补充材料 | 补充材料 |

---

## 附录：待补充信息清单

（同 v4：伦理批号、诊断标准、IQ/ADOS 量表、EEG 采集、ICA、`language_score` 等。）

---

## 附录：引用待核查清单

（同 v4；见 `reference_audit_APA.md`。）

---

## 生成统计（v5）

| 指标 | 数值 |
|------|------|
| 正文图 | 4（图 1A/1B 计为图 1 两 panel） |
| 补充图 | 2 |
| 正文表 | 2 |
| 补充表 | 6 |
"""

    out_md = body + refs_tail.split("---")[0] if "---" in refs_tail else refs_tail
    # Keep references only up to Wilkinson then add supp
    ref_match = re.search(
        r"(## References[\s\S]*?Wilkinson[^\n]+\n)",
        refs_tail,
    )
    if ref_match:
        refs_only = ref_match.group(1).strip()
        out_md = body.rstrip() + "\n\n---\n\n" + refs_only + "\n\n" + "".join(supp_parts) + appendix
    else:
        out_md = body + refs_tail + "".join(supp_parts) + appendix
    out_md = re.sub(r"\n---\n\n---\n\n## References", "\n\n---\n\n## References", out_md)

    OUT_MD.write_text(out_md, encoding="utf-8")
    print(f"Wrote {OUT_MD}")

    write_docx(out_md, audit)
    write_audit(audit)
    print("Done.")


def write_docx(
    md_text: str,
    audit: dict,
    out_path: Path = OUT_DOCX,
    title: str | None = None,
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def set_run_font(run, east_asia="宋体", latin="Times New Roman"):
        run.font.name = latin
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)

    lines = md_text.splitlines()
    i = 0
    in_refs = False
    img_re = re.compile(r"^!\[(.+?)\]\((.+?)\)$")

    while i < len(lines):
        line = lines[i]
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            h = line[3:].strip()
            if h == "References":
                in_refs = True
            elif h == "补充材料":
                in_refs = False
            doc.add_heading(h, level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        m = img_re.match(line.strip())
        if m:
            alt, rel = m.groups()
            img_path = (OUT_MD.parent / rel).resolve()
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Cm(15))
            else:
                doc.add_paragraph(f"[Figure file missing: {rel}]")
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_md_table(doc, table_lines, set_run_font)
            continue
        if line.strip().startswith("[Figure file missing"):
            doc.add_paragraph(line.strip())
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        p = doc.add_paragraph()
        add_formatted_paragraph(p, line, set_run_font)
        i += 1

    doc.save(out_path)

    if out_path == OUT_DOCX:
        print(f"Wrote {out_path}")
        supp_start = md_text.find("## 补充材料")
        if supp_start >= 0:
            write_docx(
                md_text[supp_start:],
                audit,
                OUT_SUPP_DOCX,
                title="补充材料",
            )
            print(f"Wrote {OUT_SUPP_DOCX}")


def add_md_table(doc, table_lines, set_run_font):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rows = []
    for tl in table_lines:
        if re.match(r"^\|[\s\-:|]+\|$", tl):
            continue
        cells = [c.strip() for c in tl.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            tbl.rows[ri].cells[ci].text = cell.replace("*", "")
    # three-line style: remove vertical borders (simplified)
    tbl.autofit = True


def add_formatted_paragraph(p, text, set_run_font):
    parts = re.split(r"(\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
            set_run_font(run)
        else:
            run = p.add_run(part)
            set_run_font(run)


def write_audit(audit: dict) -> None:
    lines = ["# 图表插入审计（manuscript v5）\n", f"**输出稿**：`{OUT_MD.name}`\n\n"]
    lines.append("## 1. 成功插入的图\n\n")
    lines.append("| Figure | File used | Section | Caption |\n")
    lines.append("|--------|-----------|---------|----------|\n")
    sec_map = {
        "fig1a": "3.1",
        "fig1b": "3.1",
        "fig2": "3.2",
        "fig3a": "3.4",
        "fig3b": "3.4",
        "fig4": "3.5",
        "figs1": "Supplementary",
        "figs2": "Supplementary",
    }
    for f in audit["figures"]:
        ok = "Yes" if f.get("resolved") else "No"
        rel = ""
        if f.get("resolved"):
            try:
                rel = str(Path(f["resolved"]).relative_to(PROJECT_ROOT))
            except ValueError:
                rel = f["resolved"]
        lines.append(
            f"| {f['label']} | `{rel}` | {sec_map.get(f['key'], '')} | {ok} |\n"
        )
    lines.append("\n## 2. 缺失的图\n\n")
    if audit["missing_figures"]:
        for lab, rel, checked in audit["missing_figures"]:
            lines.append(f"- **{lab}**: expected `{rel}`; checked: {checked}\n")
    else:
        lines.append("无。\n")
    lines.append("\n## 3. 成功插入的表\n\n")
    lines.append("| Table | Source | Rows | Notes |\n")
    lines.append("|-------|--------|------|-------|\n")
    lines.append("| Table 1 | table1_main_cohort_*.csv + table2_main_cohort_eeg_qc.csv | merged | 正文 §3.1 |\n")
    lines.append("| Table 2 | global_exponent_robustness_models.csv | 6 models | 正文 §3.3 |\n")
    for t in audit.get("tables", []):
        lines.append(f"| Table {t[0]} | {t[1]} | {t[2]} | {t[3]} |\n")
    lines.append("\n## 4. 缺失的表\n\n")
    if audit["missing_tables"]:
        for t in audit["missing_tables"]:
            lines.append(f"- Table {t[0]}: `{t[1]}`\n")
    else:
        lines.append("无。\n")

    # Count unused figures
    all_figs = list((PROJECT_ROOT / "outputs/figures").rglob("*.svg"))
    used = {Path(f["resolved"]).name for f in audit["figures"] if f.get("resolved")}
    unused = [p for p in all_figs if p.name not in used and "_png_cache" not in str(p)]
    lines.append("\n## 5. 未使用的候选图（部分）\n\n")
    for p in sorted(unused)[:15]:
        lines.append(f"- `{p.relative_to(PROJECT_ROOT)}`\n")

    n_body_fig = sum(1 for f in audit["figures"] if f["key"] in ("fig1a", "fig1b", "fig2", "fig3a", "fig3b", "fig4") and f.get("resolved"))
    n_supp_fig = sum(1 for f in audit["figures"] if f["key"] in ("figs1", "figs2") and f.get("resolved"))
    lines.append("\n## 汇总\n\n")
    lines.append(f"- 正文图 panel 数：{n_body_fig}\n")
    lines.append(f"- 补充图：{n_supp_fig}\n")
    lines.append("- 正文表：2\n")
    lines.append(f"- 补充表：{len(audit.get('tables', []))}\n")
    lines.append(f"- 缺失图：{len(audit['missing_figures'])}\n")
    lines.append(f"- 缺失表：{len(audit['missing_tables'])}\n")

    OUT_AUDIT.write_text("".join(lines), encoding="utf-8")
    print(f"Wrote {OUT_AUDIT}")


if __name__ == "__main__":
    main()
