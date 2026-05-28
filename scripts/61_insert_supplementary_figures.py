from pathlib import Path
import shutil
from docx import Document
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(r"d:\asd_eeg_aperiodic_study")
DOCX = ROOT / "supplementary_materials.docx"
BACKUP = ROOT / "supplementary_materials_before_figure_insertion.docx"
FIG_DIR = ROOT / "figures_submission"


FIG_MAP = {
    "Figure S1. Split-half reliability": FIG_DIR / "FigureS1_split_half_reliability.png",
    "Figure S2. Automated ICA global vs posterior results": FIG_DIR / "FigureS2_ica_global_vs_posterior.png",
    "Figure S3. Primary vs stringent sensitivity summary": FIG_DIR / "FigureS3_primary_vs_stringent.png",
    "Figure S4. Delta_Exponent results": FIG_DIR / "FigureS4_delta_exponent.png",
    "Figure S5. posterior-CARS exploratory scatter": FIG_DIR / "FigureS5_posterior_cars_scatter.png",
    "Figure S6. classification ROC/AUC/feature importance": FIG_DIR / "FigureS6_classification_summary.png",
}


def paragraph_has_drawing(p: Paragraph) -> bool:
    return bool(p._p.xpath(".//w:drawing"))


def add_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def remove_following_drawing_paragraphs(paragraph: Paragraph):
    nxt = paragraph._p.getnext()
    while nxt is not None:
        p_obj = Paragraph(nxt, paragraph._parent)
        text = p_obj.text.strip()
        if text:
            break
        if paragraph_has_drawing(p_obj):
            to_remove = nxt
            nxt = nxt.getnext()
            to_remove.getparent().remove(to_remove)
            continue
        # blank non-drawing paragraph: keep one spacer and stop
        break


def main():
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    for k, v in FIG_MAP.items():
        if not v.exists():
            raise FileNotFoundError(f"Missing figure for {k}: {v}")

    shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))

    inserted = 0
    for p in doc.paragraphs:
        title = p.text.strip()
        if title in FIG_MAP:
            remove_following_drawing_paragraphs(p)
            np = add_paragraph_after(p)
            run = np.add_run()
            run.add_picture(str(FIG_MAP[title]), width=Cm(16.0))
            inserted += 1

    doc.save(str(DOCX))
    print(f"inserted_figures={inserted}")
    print(f"inline_shapes={len(doc.inline_shapes)}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    main()
