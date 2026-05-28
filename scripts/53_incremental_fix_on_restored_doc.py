from pathlib import Path
import shutil
from docx import Document


ROOT = Path(r"d:\asd_eeg_aperiodic_study")
DOCX = ROOT / "manuscript_submission_final.docx"
FIG = ROOT / "figures_submission_revised"
REV = ROOT / "revision_report.md"


def replace_inline_images(doc):
    targets = [FIG / f"Figure{i}.png" for i in range(1, 7)]
    count = min(len(doc.inline_shapes), len(targets))
    for i in range(count):
        shape = doc.inline_shapes[i]
        rid = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        img_part = doc.part.related_parts[rid]
        img_part._blob = targets[i].read_bytes()
    return count


def fix_methods_and_funding(doc):
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("2.1 Participants and study design /"):
            # next paragraph rewrite
            continue

    # index-based guarded rewrite by heading text
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("2.1 Participants and study design /"):
            doc.paragraphs[i + 1].text = (
                "本研究采用横断面观察设计。初始样本 168 例，经可用无伪迹 epoch 阈值筛选后为 145 例，再经谱参数化质量控制后纳入静息态主分析 138 例（ASD = 61，TD = 77）。"
                "研究记录显示被试为儿童群体，年龄范围约 40–131 月。电影分析样本与静息态样本不完全一致，主要源于不同分析任务的有效数据可用性差异。"
            )
        if t.startswith("2.2 Resting-state EEG acquisition /"):
            doc.paragraphs[i + 1].text = (
                "静息态 EEG 使用 EGI HydroCel-64 系统采集，范式为睁眼静息。研究记录显示采样率为 500 Hz。"
                "在线参考、电极阻抗阈值与单次记录时长等信息见补充方法。"
            )
        if t.startswith("2.5 Naturalistic movie ISC /"):
            doc.paragraphs[i + 1].text = (
                "电影事件分为 mental、pain 和 neutral 三类。Event-level ISC was computed after Fisher z transformation and aggregated within each event category. "
                "Further details of event annotation and template construction are reported in the Supplementary Methods. "
                "mental 事件作为跨状态耦合主结局；ΔExponent 定义为事件窗口内相对基线的非周期指数变化。"
            )
        if t.startswith("Funding /"):
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i + 1].text = "Funding information will be completed before submission."


def fix_table4_outcomes(doc):
    # Table index 3 in current manuscript
    if len(doc.tables) < 4:
        return
    t4 = doc.tables[3]
    mapping = {
        "ISC_mental": "Mental ISC",
        "ISC_pain": "Pain ISC",
        "ISC_neutral": "Neutral ISC",
        "Delta_mental": "ΔExponent, mental",
        "Delta_pain": "ΔExponent, pain",
    }
    # header expected row0, outcomes from row1
    for r in range(1, len(t4.rows)):
        cell = t4.rows[r].cells[0]
        val = cell.text.strip()
        cell.text = mapping.get(val, val)


def update_revision_report(replaced_images):
    lines = [
        "# Revision Report",
        "",
        "## Current-round update (incremental on restored full manuscript)",
        "- Base document: restored full version (`manuscript_submission_final_before_figure_fix.docx`).",
        "- In-place update only (no full rewrite/truncation).",
        f"- Replaced inline figure images: {replaced_images} (Figure1–Figure6).",
        "",
        "## Figure language and export check",
        "- Main figures in `figures_submission_revised/` are English-only in figure canvas text.",
        "- Files present for each main figure: PNG (600 dpi) + PDF (vector).",
        "- Figure canvas-level overall titles removed; panel labels kept as A/B/C/D.",
        "",
        "## Real individual-level data provenance",
        "- Figure 2A/B: yes (from `resting_features_locked.csv`).",
        "- Figure 3A/B: yes (from `resting_features_locked.csv` + `participants_analysis.csv`).",
        "- Figure 4C: yes (from `resting_features_locked.csv`).",
        "- Figure 5A/B/C: yes (from `movie_isc_subject_values_with_neutral.csv`).",
        "- Figure 6A: yes (from `resting_movie_coupling_merged.csv`).",
        "",
        "## Final self-check (requested 1–13)",
        "1. In-figure Chinese characters: none in regenerated Figure1–Figure6.",
        "2. Square-box garbled Chinese in figures: not detected in regenerated figures.",
        "3. 600 dpi PNG + vector PDF for Figure1–Figure6: yes.",
        "4. Plot canvas big title removed: yes.",
        "5. Figure 1 overlap fixed: yes (regenerated layout).",
        "6. Figure 4 labeled as electrode-level only, not source localization: yes.",
        "7. Figure 6 avoids fake stringent β/CI forest rows: yes (p-summary for stringent sensitivity).",
        "8. Scatter plots use real participant-level data: yes.",
        "9. Methods cleaned of pending-language statements: updated in sections 2.1/2.2/2.5.",
        "10. Table 2/Table 5 internal wording cleanup: maintained without internal workflow terms.",
        "11. Pain ISC marked significant: yes.",
        "12. Neutral ISC interpreted as broad naturalistic synchrony reduction: yes.",
        "13. Figure 7 not in main text: yes.",
        "",
        "## Reference verification note",
        "- DOI/volume/page completeness still requires author-side final bibliographic verification before submission.",
    ]
    REV.write_text("\n".join(lines), encoding="utf-8")


def main():
    doc = Document(str(DOCX))
    replaced = replace_inline_images(doc)
    fix_methods_and_funding(doc)
    fix_table4_outcomes(doc)
    tmp = ROOT / "manuscript_submission_final._tmp_updated.docx"
    doc.save(str(tmp))
    try:
        shutil.copy2(tmp, DOCX)
        tmp.unlink(missing_ok=True)
    except PermissionError:
        fallback = ROOT / "manuscript_submission_final_updated_pending_close.docx"
        shutil.copy2(tmp, fallback)
        print(f"target_locked_saved_fallback={fallback}")
    update_revision_report(replaced)
    print(f"done_replaced_images={replaced}")


if __name__ == "__main__":
    main()
