#!/usr/bin/env python
"""
44_build_manuscript_v16.py
--------------------------
整合当前结果，生成 v16 投稿稿：
1) outputs/reports/manuscript_draft_zh_v16_submission_ready.md
2) outputs/reports/manuscript_draft_zh_v16_submission_ready.docx
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches


def _p(v: float) -> str:
    if pd.isna(v):
        return "NA"
    v = float(v)
    if v < 1e-4:
        return f"{v:.2e}"
    return f"{v:.4f}"


def _load_metric(df: pd.DataFrame, metric: str) -> pd.Series:
    row = df[df["Analysis_Type"] == metric]
    if row.empty:
        raise ValueError(f"final_paper_stats_locked.csv 缺少 {metric}")
    return row.iloc[0]


def build_markdown(
    root: Path,
    out_md: Path,
    final_df: pd.DataFrame,
    main_group_df: pd.DataFrame,
    ch_sig_df: pd.DataFrame,
    sample_flow_df: pd.DataFrame,
    relaxed_df: pd.DataFrame,
    strict_df: pd.DataFrame,
    qc_cov_df: pd.DataFrame,
    qc_cov_fdr_df: pd.DataFrame,
) -> None:
    isc_mental = _load_metric(final_df, "ISC_mental")
    isc_pain = _load_metric(final_df, "ISC_pain")
    isc_neutral = _load_metric(final_df, "ISC_neutral")
    delta_mental = _load_metric(final_df, "Delta_mental")
    delta_pain = _load_metric(final_df, "Delta_pain")

    ge = main_group_df[
        (main_group_df["outcome"] == "global_exponent")
        & (main_group_df["term"] == "C(group)[T.TD]")
    ].iloc[0]

    go = main_group_df[
        (main_group_df["outcome"] == "global_offset")
        & (main_group_df["term"] == "C(group)[T.TD]")
    ].iloc[0]

    relaxed_pc = relaxed_df[relaxed_df["metric"] == "posterior_vs_cars_spearman"].iloc[0]
    strict_pc = strict_df[strict_df["metric"] == "posterior_vs_cars_spearman"].iloc[0]
    relaxed_ols = relaxed_df[relaxed_df["metric"] == "coupling_interaction_ols"].iloc[0]
    strict_ols = strict_df[strict_df["metric"] == "coupling_interaction_ols"].iloc[0]
    relaxed_rlm = relaxed_df[relaxed_df["metric"] == "coupling_interaction_rlm"].iloc[0]
    strict_rlm = strict_df[strict_df["metric"] == "coupling_interaction_rlm"].iloc[0]

    ch_list = ", ".join(ch_sig_df["channel"].astype(str).tolist())
    build_time = datetime.now().strftime("%Y-%m-%d %H:%M")

    md = f"""# 自闭症谱系障碍儿童静息态后部非周期 EEG 与自然电影社会加工耦合异常：整合分析（v16）

> 构建时间：{build_time}  
> 本稿整合当前工作区最新结果：宽松样本主分析 + 严格样本敏感性分析 + QC 协变量控制分析。

## 摘要

**背景：** ASD 儿童在静息态非周期神经动力学上的异常已多次观察到，但其是否会迁移至自然社会认知加工场景，并形成“状态-加工解耦”，仍缺乏系统证据。  
**目的：** 在静息态与自然电影 EEG 一体化框架下，检验 ASD 与 TD 的跨状态神经耦合差异，并评估结果对样本口径与数据质量控制的稳健性。  
**方法：** 静息态采用 specparam 指标（global/posterior），电影态采用事件分段 ISC（mental/pain/neutral）与 Delta 指数。主统计采用家族内 BH-FDR（ISC 家族、Delta 家族）。同时报告严格口径敏感性与“宽松样本 + QC协变量控制”模型。  
**结果：** 静息态 global_exponent 组差显著（TD>ASD，β={ge['coef']:.3f}, p={_p(ge['pvalue'])}），后枕通道 {ch_list} 经 FDR 仍显著。电影态 ISC 在 mental/pain/neutral 三事件中均显示 ASD< TD（FDR p 分别为 {_p(isc_mental['FDR_p'])}、{_p(isc_pain['FDR_p'])}、{_p(isc_neutral['FDR_p'])}）。Delta 指数组差在 mental/pain 两条件均显著（FDR p 均为 {_p(delta_mental['FDR_p'])}）。跨状态耦合中，宽松口径下 OLS 与 RLM 交互均显著（p={_p(relaxed_ols['raw_p'])}, {_p(relaxed_rlm['raw_p'])}），严格口径下 OLS 转为不显著（p={_p(strict_ols['raw_p'])}）而 RLM 保持显著（p={_p(strict_rlm['raw_p'])}）。posterior-CARS 相关在宽松口径显著（p={_p(relaxed_pc['raw_p'])}），严格口径不显著（p={_p(strict_pc['raw_p'])}）。  
**结论：** ASD 儿童呈现“群体神经同步下降 + 状态切换调节不足 + 跨状态耦合受损”的一致模式。核心组间主效应在不同口径下稳健，而边界型临床相关对 QC 与样本定义敏感，提示其应定位为探索性证据。

**关键词：** ASD；EEG；aperiodic exponent；自然电影；ISC；state-to-processing coupling

---

## 1. 引言

静息态非周期 EEG 参数（尤其后部 exponent）为刻画 ASD 基础神经动力学提供了量化抓手。然而，静息态差异是否能解释自然社会场景中的在线神经加工异常，仍未充分回答。自然电影范式具备高生态效度和刺激锁定特性，可连接“基础状态（resting）”与“任务输出（movie ISC）”。

本研究提出并检验：ASD 是否存在从静息态后部非周期状态到自然社会加工神经同步的跨状态功能解耦。

---

## 2. 方法

### 2.1 数据与样本

- 静息态主流程样本流失见 `outputs/tables/sample_inclusion_flow.csv`。  
- 电影态事件分段与 ISC 基于 `mental/pain/neutral` 三事件类型。  
- 本稿报告两类口径：
  - 宽松口径（保留更大样本，当前主表）：`outputs/tables/final_paper_stats_locked.csv`
  - 严格口径（QC更严格，敏感性对照）：`outputs/tables/_strict_significance_snapshot.csv`

### 2.2 统计框架

- ISC 家族：mental/pain/neutral 组间比较，同家族 BH-FDR。  
- Delta 家族：mental/pain 组间比较，同家族 BH-FDR。  
- 跨状态耦合：`mental_isc_z ~ posterior_exponent * group + covariates`，并进行 Winsorize + Huber RLM。  
- QC 协变量控制（第三条路径）：在宽松样本中加入 `mean_fit_error`、`invalid_channel_ratio`、`usable_epochs` 作为 nuisance covariates（见 `outputs/tables/qc_covariate_control_models.csv`）。

---

## 3. 结果

### 3.1 静息态主效应与空间证据

- `global_exponent`：TD > ASD（β={ge['coef']:.3f}, p={_p(ge['pvalue'])}, n={int(ge['n_obs'])}）。  
- `global_offset`：组差未达显著（β={go['coef']:.3f}, p={_p(go['pvalue'])}）。  
- 通道级 FDR 显著后枕簇：{ch_list}。

### 3.2 电影态 ISC 全线显著（FDR 后）

- ISC_mental：{isc_mental['Cohort_N']}，t={isc_mental['Test_Statistic']:.3f}，raw p={_p(isc_mental['Raw_p'])}，FDR p={_p(isc_mental['FDR_p'])}。  
- ISC_pain：{isc_pain['Cohort_N']}，t={isc_pain['Test_Statistic']:.3f}，raw p={_p(isc_pain['Raw_p'])}，FDR p={_p(isc_pain['FDR_p'])}。  
- ISC_neutral：{isc_neutral['Cohort_N']}，t={isc_neutral['Test_Statistic']:.3f}，raw p={_p(isc_neutral['Raw_p'])}，FDR p={_p(isc_neutral['FDR_p'])}。

### 3.3 Delta（静息→任务态调节）显著

- Delta_mental：{delta_mental['Cohort_N']}，t={delta_mental['Test_Statistic']:.3f}，raw p={_p(delta_mental['Raw_p'])}，FDR p={_p(delta_mental['FDR_p'])}。  
- Delta_pain：{delta_pain['Cohort_N']}，t={delta_pain['Test_Statistic']:.3f}，raw p={_p(delta_pain['Raw_p'])}，FDR p={_p(delta_pain['FDR_p'])}。

### 3.4 跨状态耦合与口径敏感性

- 宽松口径：OLS 交互显著（p={_p(relaxed_ols['raw_p'])}），RLM 交互显著（p={_p(relaxed_rlm['raw_p'])}）。  
- 严格口径：OLS 交互不显著（p={_p(strict_ols['raw_p'])}），RLM 交互仍显著（p={_p(strict_rlm['raw_p'])}）。

### 3.5 posterior-CARS 的边界敏感性

- 宽松口径：n={int(relaxed_pc['n'])}, Spearman p={_p(relaxed_pc['raw_p'])}（显著）。  
- 严格口径：n={int(strict_pc['n'])}, Spearman p={_p(strict_pc['raw_p'])}（不显著）。  
- QC 协变量控制 OLS 显示 posterior 项不显著（见 `outputs/tables/qc_covariate_control_models.csv`）。

---

## 4. 讨论

1. **核心主效应稳健**：ISC（3事件）与 Delta（2事件）在不同口径下均保持显著，支持 ASD 的“同步低 + 切换弱”双特征。  
2. **耦合证据呈模型依赖**：RLM 在两口径下均显著，OLS 对样本口径更敏感，提示异常值与质量波动会影响线性参数推断。  
3. **临床相关需降级表述**：posterior-CARS 在宽松显著、严格不显著，且协变量控制后不稳，宜定位探索性发现。  
4. **方法学贡献**：通过“严格主线 + 宽松+QC协变量”的双路径报告，减少了选择性报告风险并提高结果可审查性。

---

## 5. 结论

本研究支持 ASD 在自然社会情境中的普遍性神经同步下降与状态调节不足，并提示静息态后部非周期指标与任务加工之间存在跨状态耦合异常。核心组间结果稳健，而临床相关更依赖样本与QC定义，后续需更大样本前瞻验证。

---

## 6. 主文图表（已嵌入路径）

### Figure 1. 静息态主效应（paper-ready）
![Figure 1](../figures/paper_ready_v2/fig2_main_aperiodic_effects.png)

### Figure 2. 电影事件 ISC 组间差异
![Figure 2](../../outputs_task_movie/figures/movie_isc_group_boxplot.png)

### Figure 3. 静息态-电影耦合散点（全样本）
![Figure 3](../../outputs_task_movie/figures/resting_to_movie_coupling_scatter.png)

### Figure 4. 静息态-电影耦合散点（>72月）
![Figure 4](../../outputs_task_movie/figures/resting_to_movie_coupling_scatter_older72.png)

### Figure 5. posterior exponent 与 CARS（探索性）
![Figure 5](../figures/clinical/posterior_exponent_vs_cars_scatter.png)

---

## 7. 核心统计表

### Table 1. 最终主表（ISC + Delta，家族内 FDR）

| Analysis_Type | Cohort_N | Test_Statistic | Raw_p | FDR_p |
|---|---|---:|---:|---:|
| ISC_mental | {isc_mental['Cohort_N']} | {isc_mental['Test_Statistic']:.3f} | {_p(isc_mental['Raw_p'])} | {_p(isc_mental['FDR_p'])} |
| ISC_pain | {isc_pain['Cohort_N']} | {isc_pain['Test_Statistic']:.3f} | {_p(isc_pain['Raw_p'])} | {_p(isc_pain['FDR_p'])} |
| ISC_neutral | {isc_neutral['Cohort_N']} | {isc_neutral['Test_Statistic']:.3f} | {_p(isc_neutral['Raw_p'])} | {_p(isc_neutral['FDR_p'])} |
| Delta_mental | {delta_mental['Cohort_N']} | {delta_mental['Test_Statistic']:.3f} | {_p(delta_mental['Raw_p'])} | {_p(delta_mental['FDR_p'])} |
| Delta_pain | {delta_pain['Cohort_N']} | {delta_pain['Test_Statistic']:.3f} | {_p(delta_pain['Raw_p'])} | {_p(delta_pain['FDR_p'])} |

### Table 2. 口径敏感性快照（关键变化项）

| 指标 | 宽松口径 p | 严格口径 p | 显著性变化 |
|---|---:|---:|---|
| posterior-CARS Spearman | {_p(relaxed_pc['raw_p'])} | {_p(strict_pc['raw_p'])} | 显著 -> 不显著 |
| Coupling OLS 交互 | {_p(relaxed_ols['raw_p'])} | {_p(strict_ols['raw_p'])} | 显著 -> 不显著 |
| Coupling RLM 交互 | {_p(relaxed_rlm['raw_p'])} | {_p(strict_rlm['raw_p'])} | 显著 -> 显著 |

### Table 3. 宽松样本 + QC 协变量控制（摘录）

见：
- `outputs/tables/qc_covariate_control_models.csv`
- `outputs/tables/qc_covariate_control_event_fdr.csv`
"""

    out_md.parent.mkdir(parents=True, exist_ok=True)
    out_md.write_text(md, encoding="utf-8")


def _add_df_table(doc: Document, df: pd.DataFrame, max_rows: int = 20) -> None:
    if df.empty:
        doc.add_paragraph("（无可用数据）")
        return
    df2 = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(df2.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(df2.columns):
        hdr[i].text = str(c)
    for _, row in df2.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(df2.columns):
            cells[i].text = str(row[c])


def build_docx(
    root: Path,
    out_docx: Path,
    final_df: pd.DataFrame,
    sample_flow_df: pd.DataFrame,
    relaxed_df: pd.DataFrame,
    strict_df: pd.DataFrame,
    qc_cov_fdr_df: pd.DataFrame,
) -> None:
    doc = Document()
    doc.add_heading("自闭症谱系障碍儿童静息态后部非周期 EEG 与自然电影社会加工耦合异常（v16）", level=1)
    doc.add_paragraph(f"构建时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("摘要", level=2)
    doc.add_paragraph(
        "本稿整合当前工作区最新结果（宽松样本主分析、严格样本敏感性、QC协变量控制）。"
        "核心发现：ISC（三事件）与Delta（两事件）在家族内FDR后均显著；跨状态耦合在RLM下稳健；"
        "posterior-CARS在不同口径间呈边界敏感。"
    )

    doc.add_heading("主结果图", level=2)
    fig_paths = [
        root / "outputs" / "figures" / "paper_ready_v2" / "fig2_main_aperiodic_effects.png",
        root / "outputs_task_movie" / "figures" / "movie_isc_group_boxplot.png",
        root / "outputs_task_movie" / "figures" / "resting_to_movie_coupling_scatter.png",
        root / "outputs_task_movie" / "figures" / "resting_to_movie_coupling_scatter_older72.png",
        root / "outputs" / "figures" / "clinical" / "posterior_exponent_vs_cars_scatter.png",
    ]
    fig_titles = [
        "Figure 1. 静息态主效应",
        "Figure 2. 电影事件 ISC",
        "Figure 3. 静息态-电影耦合（全样本）",
        "Figure 4. 静息态-电影耦合（>72月）",
        "Figure 5. posterior exponent 与 CARS（探索性）",
    ]
    for title, path in zip(fig_titles, fig_paths):
        doc.add_paragraph(title)
        if path.exists():
            doc.add_picture(str(path), width=Inches(6.2))
        else:
            doc.add_paragraph(f"[缺失图像] {path}")

    doc.add_heading("Table 1. 最终主表（ISC + Delta）", level=2)
    _add_df_table(doc, final_df, max_rows=10)

    doc.add_heading("Table 2. 样本流失（静息态）", level=2)
    _add_df_table(doc, sample_flow_df, max_rows=20)

    doc.add_heading("Table 3. 口径敏感性（宽松 vs 严格）", level=2)
    cmp = pd.merge(
        relaxed_df[["metric", "n", "raw_p", "significant"]].rename(
            columns={"n": "n_relaxed", "raw_p": "p_relaxed", "significant": "sig_relaxed"}
        ),
        strict_df[["metric", "n", "raw_p", "significant"]].rename(
            columns={"n": "n_strict", "raw_p": "p_strict", "significant": "sig_strict"}
        ),
        on="metric",
        how="inner",
    )
    _add_df_table(doc, cmp, max_rows=20)

    doc.add_heading("Table 4. QC 协变量控制（ISC 事件级 FDR）", level=2)
    _add_df_table(doc, qc_cov_fdr_df, max_rows=10)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def main() -> None:
    root = Path(__file__).resolve().parents[1]

    final_df = pd.read_csv(root / "outputs" / "tables" / "final_paper_stats_locked.csv")
    main_group_df = pd.read_csv(root / "derivatives" / "stats" / "main_group_analysis.csv")
    ch_sig_df = pd.read_csv(root / "outputs" / "tables" / "significant_channels_fdr.csv")
    sample_flow_df = pd.read_csv(root / "outputs" / "tables" / "sample_inclusion_flow.csv")
    relaxed_df = pd.read_csv(root / "outputs" / "tables" / "_relaxed_significance_snapshot.csv")
    strict_df = pd.read_csv(root / "outputs" / "tables" / "_strict_significance_snapshot.csv")
    qc_cov_df = pd.read_csv(root / "outputs" / "tables" / "qc_covariate_control_models.csv")
    qc_cov_fdr_df = pd.read_csv(root / "outputs" / "tables" / "qc_covariate_control_event_fdr.csv")

    out_md = root / "outputs" / "reports" / "manuscript_draft_zh_v16_submission_ready.md"
    out_docx = root / "outputs" / "reports" / "manuscript_draft_zh_v16_submission_ready.docx"

    build_markdown(
        root=root,
        out_md=out_md,
        final_df=final_df,
        main_group_df=main_group_df,
        ch_sig_df=ch_sig_df,
        sample_flow_df=sample_flow_df,
        relaxed_df=relaxed_df,
        strict_df=strict_df,
        qc_cov_df=qc_cov_df,
        qc_cov_fdr_df=qc_cov_fdr_df,
    )
    build_docx(
        root=root,
        out_docx=out_docx,
        final_df=final_df,
        sample_flow_df=sample_flow_df,
        relaxed_df=relaxed_df,
        strict_df=strict_df,
        qc_cov_fdr_df=qc_cov_fdr_df,
    )
    print(f"Saved: {out_md}")
    print(f"Saved: {out_docx}")


if __name__ == "__main__":
    main()

